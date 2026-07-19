# -*- coding: utf-8 -*-
"""WO-11 — Sinh bo 7 chung tu tu 1 bao gia + Xuat Excel/Word theo template chuan cong ty.

- sinh_bo_chung_tu(conn, sess, quotation_id): tao 7 chung tu NHAP da dien san (khong tu chot).
- export_doc(conn, loai, id, fmt): tra (filename, bytes) — dien {{placeholder}} vao template
  chuan tai D:\\Quản trị DOANH NGHIỆP\\Mẫu chứng từ chuẩn\\; thieu template -> dung ban sach.
"""
import glob
import io
import os
import re
import sys
import unicodedata
import uuid
import zipfile
from datetime import date, datetime
from decimal import Decimal, InvalidOperation

import db as D
from api_write import (ValidationError, audit, next_code, require_write,
                       luu_file_vao_folder_khach, dam_bao_folder_khach)

try:
    sys.stdout.reconfigure(encoding="utf-8")
except Exception:
    pass

TPL_ROOT = r"D:\Quản trị DOANH NGHIỆP\Mẫu chứng từ chuẩn"
TPL_GLOB = {
    ("bbnt", "xlsx"): "02*/03_TEMPLATE_BIEN_BAN_NGHIEM_THU*.xlsx",
    ("bbnt", "docx"): "02*/03_TEMPLATE_BIEN_BAN_NGHIEM_THU*.docx",
    ("bqt", "xlsx"): "03*/04_TEMPLATE_BANG_QUYET_TOAN*.xlsx",
    ("bqt", "docx"): "03*/04_TEMPLATE_BANG_QUYET_TOAN*.docx",
    ("dccn", "xlsx"): "04*/06_TEMPLATE_BIEN_BAN_DOI_CHIEU*.xlsx",
    ("dccn", "docx"): "04*/06_TEMPLATE_BIEN_BAN_DOI_CHIEU*.docx",
    ("payment", "xlsx"): "05*/05_TEMPLATE_THU_DE_NGHI_THANH_TOAN*.xlsx",
    ("payment", "docx"): "05*/05_TEMPLATE_THU_DE_NGHI_THANH_TOAN*.docx",
    ("pxk", "xlsx"): "07*/07_TEMPLATE_PHIEU_GIAO_HANG_XUAT_KHO*.xlsx",
    ("pxk", "docx"): "07*/07_TEMPLATE_PHIEU_GIAO_HANG_XUAT_KHO*.docx",
    ("checklist", "xlsx"): "08*/08_TEMPLATE_CHECKLIST*.xlsx",
    ("checklist", "docx"): "08*/08_TEMPLATE_CHECKLIST*.docx",
    ("quotation", "xlsx"): "01*/THANH_HOAI_TEMPLATE_BAO_GIA*.xlsx",
    ("hop_dong", "docx"): "06*/03_*THI_CONG_LAP_DAT*v2*.docx",
}


def _find_tpl(loai, fmt):
    pat = TPL_GLOB.get((loai, fmt))
    if not pat:
        return None
    hits = glob.glob(os.path.join(TPL_ROOT, pat))
    return hits[0] if hits else None


def fmt_vnd(v):
    try:
        return "{:,.0f}".format(float(v or 0)).replace(",", ".")
    except (TypeError, ValueError):
        return "0"


def fmt_d(s):
    if not s:
        return ""
    s = str(s)[:10]
    p = s.split("-")
    return "%s/%s/%s" % (p[2], p[1], p[0]) if len(p) == 3 else s


# ==================== SINH BO 7 CHUNG TU ==================================
def sinh_bo_chung_tu(conn, sess, quotation_id):
    """Tao 7 chung tu NHAP tu bao gia (khong trung: bao da co -> tra ve bo cu)."""
    require_write("sinh_chung_tu", sess["role"])
    q = conn.execute("""SELECT q.*, c.customer_name FROM quotation q
                        JOIN customer c ON c.id=q.customer_id WHERE q.id=?""",
                     (quotation_id,)).fetchone()
    if not q:
        raise ValidationError("Báo giá không tồn tại.")
    items = conn.execute("SELECT * FROM quotation_item WHERE quotation_id=? ORDER BY stt",
                         (quotation_id,)).fetchall()
    cid, pid = q["customer_id"], q["project_id"]
    ket_qua, da_co = {}, []
    try:
        dam_bao_folder_khach(conn, cid)  # dam bao co san folder khach truoc khi sinh bo
    except Exception:
        pass

    # 1) HOP DONG
    row = conn.execute("SELECT id, code FROM hop_dong_ct WHERE quotation_id=?", (quotation_id,)).fetchone()
    if row:
        da_co.append("hop_dong")
        ket_qua["hop_dong"] = dict(row)
    else:
        code = next_code(conn, "hop_dong_ct", "HD")
        conn.execute("""INSERT INTO hop_dong_ct(code, customer_id, quotation_id, loai_hd, gia_tri,
                        trang_thai) VALUES(?,?,?,?,?,?)""",
                     (code, cid, quotation_id, "Thi công lắp đặt", q["grand_total"], "Nhap"))
        ket_qua["hop_dong"] = {"id": conn.execute("SELECT last_insert_rowid()").fetchone()[0], "code": code}

    # 2) BBNT — keo hang muc tu bao gia
    row = conn.execute("SELECT id, code FROM bbnt WHERE project_id IS ? AND customer_id=? AND code LIKE 'NT-%' AND id IN (SELECT id FROM bbnt WHERE customer_id=?) AND EXISTS (SELECT 1 FROM bbnt_item bi WHERE bi.bbnt_id=bbnt.id)", (pid, cid, cid)).fetchone() if False else None
    row = conn.execute("""SELECT b.id, b.code FROM bbnt b
                          JOIN activity_log a ON a.ref_code=b.code AND a.mo_ta LIKE '%' || ? || '%'
                          LIMIT 1""", (q["code"],)).fetchone()
    if row:
        da_co.append("bbnt")
        ket_qua["bbnt"] = dict(row)
    else:
        code = next_code(conn, "bbnt", "NT")
        conn.execute("""INSERT INTO bbnt(code, customer_id, project_id, ngay_nghiem_thu,
                        dai_dien_b, ket_luan, trang_thai)
                        VALUES(?,?,?,?,?,?,?)""",
                     (code, cid, pid, date.today().isoformat(), "Đại diện Thanh Hoài", "Đạt", "Nhap"))
        bid = conn.execute("SELECT last_insert_rowid()").fetchone()[0]
        for it in items:
            conn.execute("""INSERT INTO bbnt_item(bbnt_id, hang_muc, don_gia, thanh_tien,
                            kl_hop_dong, kl_thuc_te, ket_qua) VALUES(?,?,?,?,?,?,?)""",
                         (bid, it["hang_muc"], it["don_gia"], it["thanh_tien"],
                          it["khoi_luong"], it["khoi_luong"], "Đạt"))
        conn.execute("""INSERT INTO activity_log(customer_id, project_id, loai, ref_code, mo_ta)
                        VALUES(?,?,?,?,?)""",
                     (cid, pid, "BBNT", code, "Sinh tu bao gia %s" % q["code"]))
        ket_qua["bbnt"] = {"id": bid, "code": code}

    # 3) CHECKLIST
    row = conn.execute("SELECT id, code FROM checklist_ct WHERE quotation_id=?", (quotation_id,)).fetchone()
    if row:
        da_co.append("checklist")
        ket_qua["checklist"] = dict(row)
    else:
        code = next_code(conn, "checklist_ct", "CL")
        conn.execute("""INSERT INTO checklist_ct(code, customer_id, quotation_id, loai_viec, trang_thai)
                        VALUES(?,?,?,?,?)""",
                     (code, cid, quotation_id, q["nhom_dich_vu"] or "Lắp đặt", "Nhap"))
        clid = conn.execute("SELECT last_insert_rowid()").fetchone()[0]
        for hm in ["Vệ sinh khu vực thi công", "Kiểm tra vận hành thiết bị", "Đo thông số (gas, dòng điện)",
                   "Ảnh trước / sau", "Bàn giao hướng dẫn sử dụng", "Thu dọn vật tư thừa"]:
            conn.execute("INSERT INTO checklist_dong(checklist_id, hang_muc, bat_buoc) VALUES(?,?,1)",
                         (clid, hm))
        ket_qua["checklist"] = {"id": clid, "code": code}

    # 4) BQT — cot bao gia dien san, cot thuc te de nguoi dung sua
    row = conn.execute("""SELECT id, code FROM bqt WHERE customer_id=? AND project_id IS ?
                          ORDER BY id DESC LIMIT 1""", (cid, pid)).fetchone()
    exist_for_quote = conn.execute("""SELECT b.id, b.code FROM bqt b JOIN activity_log a
                          ON a.ref_code=b.code AND a.mo_ta LIKE '%' || ? || '%' LIMIT 1""",
                                   (q["code"],)).fetchone()
    if exist_for_quote:
        da_co.append("bqt")
        ket_qua["bqt"] = dict(exist_for_quote)
    else:
        code = next_code(conn, "bqt", "BQT")
        conn.execute("""INSERT INTO bqt(code, customer_id, project_id, gia_tri_quyet_toan, da_thu,
                        con_lai, ngay_lap, trang_thai) VALUES(?,?,?,?,?,?,?,?)""",
                     (code, cid, pid, q["grand_total"], 0, q["grand_total"],
                      date.today().isoformat(), "Nhap"))
        bqid = conn.execute("SELECT last_insert_rowid()").fetchone()[0]
        for it in items:
            conn.execute("""INSERT INTO bqt_item(bqt_id, hang_muc, bao_gia, hop_dong, thuc_te,
                            phat_sinh, don_gia, thanh_tien) VALUES(?,?,?,?,?,?,?,?)""",
                         (bqid, it["hang_muc"], it["khoi_luong"], it["khoi_luong"],
                          it["khoi_luong"], "0", it["don_gia"], it["thanh_tien"]))
        conn.execute("""INSERT INTO activity_log(customer_id, project_id, loai, ref_code, mo_ta)
                        VALUES(?,?,?,?,?)""",
                     (cid, pid, "BQT", code, "Sinh tu bao gia %s" % q["code"]))
        ket_qua["bqt"] = {"id": bqid, "code": code}

    # 5) PXK — chi lay dong co ve vat tu (heuristic: khong phai nhan cong)
    row = conn.execute("SELECT id, code FROM pxk WHERE quotation_id=?", (quotation_id,)).fetchone()
    if row:
        da_co.append("pxk")
        ket_qua["pxk"] = dict(row)
    else:
        code = next_code(conn, "pxk", "PXK")
        conn.execute("""INSERT INTO pxk(code, customer_id, quotation_id, ngay_xuat, trang_thai)
                        VALUES(?,?,?,?,?)""",
                     (code, cid, quotation_id, date.today().isoformat(), "Nhap"))
        pxid = conn.execute("SELECT last_insert_rowid()").fetchone()[0]
        for it in items:
            if re.search(r"nhân công|nhan cong|lắp đặt trọn gói", it["hang_muc"], re.I):
                continue
            conn.execute("INSERT INTO pxk_dong(pxk_id, ten_hang, dvt, so_luong) VALUES(?,?,?,?)",
                         (pxid, it["hang_muc"], "", 1))
        ket_qua["pxk"] = {"id": pxid, "code": code}

    # 6) THU DE NGHI THANH TOAN
    row = conn.execute("""SELECT id, code FROM payment_request WHERE customer_id=? AND bqt_id=?
                          LIMIT 1""", (cid, ket_qua["bqt"]["id"])).fetchone()
    if row:
        da_co.append("payment")
        ket_qua["payment"] = dict(row)
    else:
        code = next_code(conn, "payment_request", "PR")
        conn.execute("""INSERT INTO payment_request(code, customer_id, bqt_id, project_id,
                        dot_thanh_toan, grand_total, status) VALUES(?,?,?,?,?,?,?)""",
                     (code, cid, ket_qua["bqt"]["id"], pid, "Quyết toán", q["grand_total"], "Nhap"))
        ket_qua["payment"] = {"id": conn.execute("SELECT last_insert_rowid()").fetchone()[0], "code": code}

    # 7) DCCN
    ky = date.today().strftime("%m/%Y")
    row = conn.execute("SELECT id, code FROM dccn WHERE customer_id=? AND ky=?", (cid, ky)).fetchone()
    if row:
        da_co.append("dccn")
        ket_qua["dccn"] = dict(row)
    else:
        code = next_code(conn, "dccn", "DCCN")
        con_no = conn.execute("""SELECT COALESCE(SUM(tong_cong - da_thu),0) FROM hoa_don
                                 WHERE customer_id=? AND chieu='ban_ra'""", (cid,)).fetchone()[0]
        conn.execute("""INSERT INTO dccn(code, customer_id, ky, du_dau, phat_sinh_tang, da_thu,
                        du_cuoi, chenh_lech, trang_thai) VALUES(?,?,?,?,?,?,?,?,?)""",
                     (code, cid, ky, 0, q["grand_total"], 0, con_no, 0, "Nhap"))
        ket_qua["dccn"] = {"id": conn.execute("SELECT last_insert_rowid()").fetchone()[0], "code": code}

    audit(conn, sess, "sinh_bo", "quotation", quotation_id,
          "Sinh bo chung tu tu %s (da co: %s)" % (q["code"], ",".join(da_co) or "khong"))
    conn.commit()

    # XUAT LUON 7 file ra dung thu muc con cua khach — khach mong doi bam "Sinh bo"
    # la thay FILE trong folder ngay, khong phai bam "Xuat" tung cai. Moi loai xuat o
    # dinh dang phu hop: van ban/thu -> Word; bang bieu -> Excel. Best-effort tung cai.
    FMT_BO = {"hop_dong": "docx", "bbnt": "docx", "checklist": "docx",
              "payment": "docx", "dccn": "docx", "bqt": "xlsx", "pxk": "xlsx"}
    da_xuat, loi_xuat = [], []
    for loai, info in ket_qua.items():
        try:
            fname, _ = export_doc(conn, loai, info["id"], FMT_BO.get(loai, "docx"))
            da_xuat.append(fname)
        except Exception as e:
            loi_xuat.append("%s: %s" % (loai, e))
    fo = conn.execute("SELECT duong_dan_folder FROM customer WHERE id=?", (cid,)).fetchone()
    folder = fo["duong_dan_folder"] if fo else None
    return {"bao_gia": q["code"], "chung_tu": ket_qua, "da_co_truoc": da_co,
            "da_xuat_file": da_xuat, "so_file": len(da_xuat), "folder": folder,
            "loi_xuat": loi_xuat}


# ==================== EXPORT XLSX / DOCX ==================================
def _data_map(conn, loai, doc, items, cfg):
    """Map {{placeholder}} -> gia tri. Alias nhieu ten cho cung 1 y."""
    kh = conn.execute("SELECT * FROM customer WHERE id=?", (doc["customer_id"],)).fetchone() if doc["customer_id"] else None
    today = date.today()
    m = {
        # cong ty
        "company_name": cfg.get("ten_cong_ty") or "", "company_tax_id": cfg.get("ma_so_thue") or "",
        "company_address": cfg.get("dia_chi") or "", "company_phone": cfg.get("dien_thoai") or "",
        "company_email": "", "bank_account_no": "", "bank_name": "",
        # khach
        "customer_name": kh["customer_name"] if kh else "", "party_a_name": kh["customer_name"] if kh else "",
        "customer_tax_id": (kh["tax_id"] or "") if kh else "", "customer_address": (kh["dia_chi"] or "") if kh else "",
        "party_a_rep": (kh["nguoi_lien_he"] or "") if kh else "", "party_a_title": "",
        "party_b_name": cfg.get("ten_cong_ty") or "", "party_b_rep": "", "party_b_title": "Giám đốc",
        # ngay
        "day": "%02d" % today.day, "month": "%02d" % today.month, "year": str(today.year),
        "date": fmt_d(today.isoformat()),
    }
    code = doc["code"] if "code" in doc.keys() else ""
    if loai == "bbnt":
        m.update({"bbnt_no": code, "acceptance_date": fmt_d(doc["ngay_nghiem_thu"]),
                  "project_name": "", "project_address": doc["dia_diem"] or "",
                  "acceptance_type": "Nghiệm thu hoàn thành", "acceptance_subject": "Nghiệm thu hoàn thành hạng mục",
                  "start_time": "", "end_time": "", "start_date": "", "end_date": "",
                  "conclusion": doc["ket_luan"] or "", "pending_items": doc["ton_dong"] or "Không",
                  "warranty": doc["thoi_han_bao_hanh"] or ""})
    elif loai == "bqt":
        m.update({"bqt_no": code, "settlement_no": code, "settlement_date": fmt_d(doc["ngay_lap"]),
                  "total_settlement": fmt_vnd(doc["gia_tri_quyet_toan"]),
                  "paid_amount": fmt_vnd(doc["da_thu"]), "remaining_amount": fmt_vnd(doc["con_lai"]),
                  # WO-29fix: template BQT chuan dung 2 token nay — truoc day bi blank
                  "finalization_date": fmt_d(doc["ngay_lap"]),
                  "remaining_in_words": so_thanh_chu(doc["con_lai"])})
    elif loai == "payment":
        m.update({"payment_no": code, "request_no": code,
                  "amount": fmt_vnd(doc["grand_total"]), "total_amount": fmt_vnd(doc["grand_total"]),
                  "payment_stage": doc["dot_thanh_toan"] or "", "due_date": fmt_d(doc["han_thanh_toan"])})
    elif loai == "dccn":
        m.update({"dccn_no": code, "period": doc["ky"] or "",
                  "opening_balance": fmt_vnd(doc["du_dau"]), "increase": fmt_vnd(doc["phat_sinh_tang"]),
                  "paid": fmt_vnd(doc["da_thu"]), "closing_balance": fmt_vnd(doc["du_cuoi"]),
                  "difference": fmt_vnd(doc["chenh_lech"])})
    elif loai == "pxk":
        m.update({"pxk_no": code, "export_date": fmt_d(doc["ngay_xuat"]),
                  "warehouse": doc["kho"] or "", "receiver": doc["nguoi_nhan"] or "",
                  # WO-29fix: alias dung TEN TOKEN THAT cua template PXK chuan (truoc bi blank)
                  "delivery_note_no": code, "posting_date": fmt_d(doc["ngay_xuat"]),
                  "receiver_name": doc["nguoi_nhan"] or ""})
    elif loai == "checklist":
        m.update({"checklist_no": code, "work_type": doc["loai_viec"] or ""})
    elif loai == "quotation":
        cols = doc.keys()
        truoc = doc["tong_truoc_thue"] if "tong_truoc_thue" in cols and doc["tong_truoc_thue"] is not None else doc["grand_total"]
        thue = doc["tien_thue"] if "tien_thue" in cols and doc["tien_thue"] is not None else 0
        m.update({"quotation_no": code, "quotation_date": fmt_d(doc["ngay_lap"]),
                  "total": fmt_vnd(doc["grand_total"]), "grand_total": fmt_vnd(doc["grand_total"]),
                  "sub_total": fmt_vnd(truoc), "vat": fmt_vnd(thue), "vat_amount": fmt_vnd(thue),
                  "service_group": doc["nhom_dich_vu"] or "",
                  # placeholder rieng cua template bao gia chuan cong ty
                  "quote_no": code, "quotation_id": code, "quote_date": fmt_d(doc["ngay_lap"]),
                  "quote_type": doc["nhom_dich_vu"] or "Báo giá",
                  "project_name": "", "valid_days": "30",
                  "payment_terms": "Theo thỏa thuận hợp đồng",
                  "template_issue_date": fmt_d(date.today().isoformat()),
                  "amount_in_words": so_thanh_chu(doc["grand_total"])})
    elif loai == "hop_dong":
        m.update({"contract_no": code, "contract_value": fmt_vnd(doc["gia_tri"]),
                  "sign_date": fmt_d(doc["ngay_ky"]) or fmt_d(today.isoformat())})
    return m


PLACEHOLDER_RE = re.compile(r"\{\{\s*([a-zA-Z0-9_]+)\s*\}\}")
# WO-34A: luoi quet CUOI — moi {{...}} con sot sau khi dien (ba cham danh dau o dien tay,
# ghi chu kieu {{DẤU NGOẶC NHỌN}} co dau/khoang trang...) -> xoa sach ("khong con {{ }} thua",
# DoD WO34A). Template cu chi dung token ASCII_underscore nen khong bi anh huong khac.
_LEFTOVER_PH_RE = re.compile(r"\{\{[^{}]*\}\}")


def _fmt_qty(v):
    """Render numeric quantities without binary-float or six-digit rounding."""
    if v in (None, ""):
        return ""
    try:
        number = Decimal(str(v))
        if not number.is_finite():
            return str(v)
        rendered = format(number, "f")
        if "." in rendered:
            rendered = rendered.rstrip("0").rstrip(".")
        return rendered or "0"
    except (InvalidOperation, TypeError, ValueError):
        return str(v)


def _item_field(it, base):
    """WO-29fix: gia tri 1 cot dong hang muc theo TEN TOKEN (khong phan biet bang nguon —
    quotation_item/bbnt_item/pxk_dong/bqt_item/dong synthetic DCCN deu tra dung cot).
    Template chuan cong ty dung token DANH SO ({{item_name_1}}, {{uom_2}}...) — day la
    resolver chung cho ca docx lan xlsx."""
    def g(*keys):
        for k in keys:
            v = it.get(k)
            if v not in (None, ""):
                return v
        return ""
    if base in ("item_name", "quotation_item", "item", "hang_muc", "ten_hang",
                "description", "giai_doan"):
        return g("hang_muc", "ten_hang", "description")
    if base in ("spec", "model", "quy_cach"):
        return g("quy_cach_model", "model", "spec", "quy_cach")
    if base in ("uom", "dvt"):
        return g("dvt") or _parse_kl(g("khoi_luong") or g("kl_hop_dong"))[1]
    if base in ("qty", "quantity", "so_luong"):
        v = it.get("so_luong")
        return _fmt_qty(v if v not in (None, "") else _parse_kl(g("khoi_luong"))[0])
    if base == "quoted_qty":
        v = g("kl_hop_dong", "bao_gia", "khoi_luong")
        return _fmt_qty(_parse_kl(v)[0]) if v else ""
    if base == "actual_qty":
        v = g("kl_thuc_te", "thuc_te", "khoi_luong")
        return _fmt_qty(_parse_kl(v)[0]) if v else ""
    if base in ("don_gia", "unit_price", "price"):
        return fmt_vnd(it["don_gia"]) if it.get("don_gia") is not None else ""
    if base in ("amount", "thanh_tien"):
        return fmt_vnd(it["thanh_tien"]) if it.get("thanh_tien") is not None else ""
    if base in ("note", "ghi_chu"):
        return g("ghi_chu", "note")
    if base in ("result", "ket_qua"):
        return g("ket_qua")
    if base == "date":                       # dong doi chieu DCCN (synthetic)
        return g("date", "ngay")
    if base in ("increase", "payment", "balance"):
        return g(base)
    return ""


def _fill_text(text, m, items, item_cursor):
    """Thay {{token}}; token khong biet -> chuoi rong."""
    def rep(match):
        key = match.group(1)
        if key in m:
            return str(m[key])
        # WO-29fix: token DANH SO {{<base>_N}} -> items[N-1] (template chuan cong ty
        # dung kieu nay; truoc day bi thay bang chuoi rong -> BBNT xuat ra trong ruot)
        base, _, num = key.rpartition("_")
        if base and num.isdigit():
            n = int(num)
            return str(_item_field(items[n - 1], base)) if 0 < n <= len(items) else ""
        # token dong hang muc kieu con tro (khong danh so)
        if key in ("quotation_item", "item_name", "model", "spec", "item"):
            i = item_cursor[0]
            if i < len(items):
                it = items[i]
                if key in ("quotation_item", "item_name", "item", "model"):
                    return str(it.get("hang_muc") or it.get("ten_hang") or "")
                return str(it.get("khoi_luong") or "")
        return ""
    return _LEFTOVER_PH_RE.sub("", PLACEHOLDER_RE.sub(rep, text))


_CHU_SO = ["không", "một", "hai", "ba", "bốn", "năm", "sáu", "bảy", "tám", "chín"]


def so_thanh_chu(n):
    """Doc so tien thanh chu tieng Viet (du dung cho bao gia)."""
    n = int(round(float(n or 0)))
    if n == 0:
        return "Không đồng"

    def doc_3(so, day_du):
        tram, chuc, dv = so // 100, (so % 100) // 10, so % 10
        p = []
        if tram or day_du:
            p += [_CHU_SO[tram], "trăm"]
        if chuc > 1:
            p.append(_CHU_SO[chuc] + " mươi")
            if dv == 1:
                p.append("mốt")
            elif dv == 5:
                p.append("lăm")
            elif dv:
                p.append(_CHU_SO[dv])
        elif chuc == 1:
            p.append("mười")
            if dv == 5:
                p.append("lăm")
            elif dv:
                p.append(_CHU_SO[dv])
        elif dv:
            if tram or day_du:
                p.append("lẻ")
            p.append(_CHU_SO[dv])
        return " ".join(p)

    ty, trieu, nghin, le = n // 10**9, (n % 10**9) // 10**6, (n % 10**6) // 1000, n % 1000
    out = []
    if ty:
        out += [doc_3(ty, False), "tỷ"]
    if trieu:
        out += [doc_3(trieu, bool(ty)), "triệu"]
    if nghin:
        out += [doc_3(nghin, bool(ty or trieu)), "nghìn"]
    if le:
        out.append(doc_3(le, bool(ty or trieu or nghin)))
    s = " ".join(out) + " đồng"
    return s[0].upper() + s[1:]


def _parse_kl(khoi_luong):
    """'10.0 m' -> (10.0, 'm'); '2 Bộ' -> (2, 'Bộ'); loi -> (1, nguyen van)."""
    parts = str(khoi_luong or "").strip().split(None, 1)
    try:
        sl = float(parts[0])
        return sl, (parts[1] if len(parts) > 1 else "")
    except (ValueError, IndexError):
        return 1, str(khoi_luong or "")


_XLREF_RE = re.compile(r"(\$?[A-Z]{1,3}\$?)(\d+)")


def _mo_rong_khu(ws, insert_at, extra, style_row):
    """WO-29fix2: chen `extra` dong BEN TRONG 1 khu hang muc cua template bao gia
    (chen truoc dong cuoi khu de SUM(H22:H31) tu gian thanh SUM(H22:H31+extra)).
    openpyxl insert_rows KHONG tu dich cong thuc + merged cells -> tu dich ca hai:
    moi tham chieu/vung merge co dong >= insert_at deu +extra (dong < giu nguyen).
    Dong moi chen: copy style + chieu cao tu style_row + cong thuc THANH TIEN tung dong."""
    from copy import copy
    if extra <= 0:
        return
    ws.insert_rows(insert_at, extra)
    # 1) dich merged cells nam duoi diem chen (shift range TRUC TIEP — khong unmerge/merge
    # vi placeholder MergedCell da bi insert_rows doi cho, unmerge se KeyError)
    for rng in ws.merged_cells.ranges:
        if rng.min_row >= insert_at:
            rng.shift(0, extra)
    # 2) dich cong thuc (chuoi '=' khong duoc openpyxl tu cap nhat)
    def _shift(f):
        return _XLREF_RE.sub(lambda mm: mm.group(1) + str(int(mm.group(2)) + (
            extra if int(mm.group(2)) >= insert_at else 0)), f)
    for row in ws.iter_rows():
        for c in row:
            if isinstance(c.value, str) and c.value.startswith("="):
                c.value = _shift(c.value)
    # 3) dong moi: style + chieu cao + cong thuc thanh tien giong dong mau
    for r in range(insert_at, insert_at + extra):
        if ws.row_dimensions[style_row].height:
            ws.row_dimensions[r].height = ws.row_dimensions[style_row].height
        for col in range(1, 10):
            ws.cell(row=r, column=col)._style = copy(ws.cell(row=style_row, column=col)._style)
        ws.cell(row=r, column=8,
                value='=IFERROR(E{r}*F{r}*(1-IF(G{r}="",0,G{r})),0)'.format(r=r))


def _fill_quote_template(tpl_path, m, items):
    """Dien template bao gia chuan cong ty: khu VAT 10% (dong 22-31) + khu VAT khac (35-44),
    cong thuc H46-H54 tu tinh. WO-29fix2: bao gia NHIEU HON 10 dong/khu -> TU GIAN KHU
    (chen dong + dich cong thuc/merge), khong con roi ve ban sach (bug chu bao 2026-07-10,
    BG-2026-0096 co 226 dong). Chi con fallback khi tron >1 thue suat ngoai 10%."""
    import openpyxl

    def vat_of(it):
        v = it.get("thue_suat")
        return 10.0 if v is None else float(v)  # LUU Y: 0% la gia tri hop le, khong duoc coi la thieu
    s10 = [it for it in items if vat_of(it) == 10]
    khac = [it for it in items if vat_of(it) != 10]
    rates = {vat_of(it) for it in khac}
    if len(rates) > 1:
        return None  # 2+ thue suat ngoai 10% — bo cuc 2 khu khong bieu dien duoc
    wb = openpyxl.load_workbook(tpl_path)
    ws = wb.active
    extra_a = max(0, len(s10) - 10)
    extra_b = max(0, len(khac) - 10)
    _mo_rong_khu(ws, 31, extra_a, style_row=22)                 # khu VAT10: 22..31+extra_a
    _mo_rong_khu(ws, 44 + extra_a, extra_b, style_row=35 + extra_a)  # khu B da bi dich xuong
    for row in ws.iter_rows():
        for c in row:
            if isinstance(c.value, str) and "{{" in c.value:
                c.value = _fill_text(c.value, m, [], [0])
    def put(row0, its):
        for i, it in enumerate(its):
            r = row0 + i
            sl, dvt = _parse_kl(it.get("khoi_luong"))
            ws.cell(row=r, column=1, value=i + 1)
            ws.cell(row=r, column=3, value=it.get("hang_muc") or it.get("ten_hang") or "")
            ws.cell(row=r, column=4, value=dvt)
            ws.cell(row=r, column=5, value=sl)
            ws.cell(row=r, column=6, value=float(it.get("don_gia") or 0))
    put(22, s10)
    put(35 + extra_a, khac)
    if khac and rates:
        off = extra_a + extra_b
        ws.cell(row=51 + off, column=8, value=list(rates)[0] / 100.0)  # thue suat thuc khu B
        ws.cell(row=50 + off, column=6, value="Subtotal (VAT %g%%)" % list(rates)[0])
        ws.cell(row=53 + off, column=6, value="Total VAT %g%%" % list(rates)[0])
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _xml_escape(s):
    return (str(s).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            .replace('"', "&quot;"))


_ITEM_PH_RE = re.compile(r"\{\{\s*(quotation_item|item_name|model|spec|item)\s*\}\}")
_ROW_RE = re.compile(r"<row\b[^>]*>.*?</row>", re.S)


def _fill_xlsx_ziplevel(tpl_path, m, items):
    """Dien placeholder o MUC ZIP — KHONG cho openpyxl mo-luu lai (tranh lam hong
    template co metadata/cong thuc hien dai, vd BQT). Giu nguyen 100% template,
    chi doi chu {{...}} trong sharedStrings.xml + sheet.
    - Vo huong: {{key}} (key trong m) -> gia tri.
    - Dong hang muc: {{quotation_item}}... theo con tro tang moi dong co placeholder."""
    zin = zipfile.ZipFile(tpl_path)
    buf = io.BytesIO()
    zout = zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED)

    def rep_scalar(match):
        key = match.group(1)
        if key in m:
            return _xml_escape(m[key])
        # WO-29fix: token danh so {{<base>_N}} trong xlsx (dong bo voi _fill_text ben docx)
        base, _, num = key.rpartition("_")
        if base and num.isdigit():
            n = int(num)
            return _xml_escape(_item_field(items[n - 1], base)) if 0 < n <= len(items) else ""
        return ""

    def fill_items_by_row(sxml):
        cursor = [0]

        def do_row(rm):
            row_xml = rm.group(0)
            if _ITEM_PH_RE.search(row_xml):
                idx = cursor[0]

                def rep_item(im):
                    key = im.group(1)
                    if idx < len(items):
                        it = items[idx]
                        if key in ("quotation_item", "item_name", "item", "model"):
                            return _xml_escape(it.get("hang_muc") or it.get("ten_hang") or "")
                        return _xml_escape(it.get("khoi_luong") or "")
                    return ""
                row_xml = _ITEM_PH_RE.sub(rep_item, row_xml)
                cursor[0] += 1
            return row_xml
        return _ROW_RE.sub(do_row, sxml)

    for info in zin.infolist():
        data = zin.read(info.filename)
        fn = info.filename
        if fn == "xl/sharedStrings.xml" or re.search(r"worksheets/sheet\d+\.xml$", fn):
            txt = data.decode("utf-8")
            if fn.endswith(".xml") and "worksheets/sheet" in fn:
                txt = fill_items_by_row(txt)          # dong hang muc (neu co)
            txt = PLACEHOLDER_RE.sub(rep_scalar, txt)  # vo huong
            data = txt.encode("utf-8")
        zout.writestr(info, data)
    zout.close()
    return buf.getvalue()


def _export_xlsx(tpl_path, m, items):
    import openpyxl
    if tpl_path and os.path.isfile(tpl_path):
        # Zip-level (khong round-trip openpyxl) — giu template nguyen ven, tranh hong file
        return _fill_xlsx_ziplevel(tpl_path, m, items)
    if True:
        # fallback: ban sach chuan cong ty (khong co template -> dung openpyxl build moi)
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.append([m.get("company_name", "")])
        ws.append(["MST: " + m.get("company_tax_id", ""), "", m.get("company_address", "")])
        ws.append([])
        ws.append([m.get("_title", "CHỨNG TỪ")])
        ws.append(["Khách hàng:", m.get("customer_name", ""), "Ngày:", m.get("date", "")])
        ws.append([])
        co_vat = any(it.get("thue_suat") is not None for it in items)
        header = ["STT", "Hạng mục", "Khối lượng", "Đơn giá", "Thành tiền"] + (["Thuế %", "Tiền thuế"] if co_vat else [])
        ws.append(header)
        for i, it in enumerate(items, 1):
            row = [i, it.get("hang_muc") or it.get("ten_hang") or "", it.get("khoi_luong") or "",
                   it.get("don_gia") or 0, it.get("thanh_tien") or 0]
            if co_vat:
                row += [it.get("thue_suat"), it.get("tien_thue") or 0]
            ws.append(row)
        ws.append([])
        truoc = sum(float(it.get("thanh_tien") or 0) for it in items)
        if co_vat:
            thue = sum(float(it.get("tien_thue") or 0) for it in items)
            ws.append(["", "", "", "Cộng trước thuế", truoc])
            ws.append(["", "", "", "Tiền thuế (VAT)", thue])
            ws.append(["", "", "", "TỔNG CỘNG", truoc + thue])
        else:
            ws.append(["", "", "", "TỔNG", truoc])
        ws.append([])
        ws.append(["ĐẠI DIỆN BÊN A", "", "", "ĐẠI DIỆN BÊN B"])
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


_IDX_TOKEN_RE = re.compile(r"\{\{\s*([a-zA-Z_]+?)_(\d+)\s*\}\}")


def _normalise_docx_header(value):
    """Stable, accent-insensitive key used only to locate a template table."""
    text = str(value or "").strip().lower().replace("đ", "d")
    text = unicodedata.normalize("NFD", text)
    text = "".join(ch for ch in text if unicodedata.category(ch) != "Mn")
    return re.sub(r"[^a-z0-9]+", " ", text).strip()


def _set_docx_cell_text_preserve_style(cell, value):
    """Replace a prototype cell's text without replacing its paragraphs/style XML."""
    first_run = None
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            if first_run is None:
                first_run = run
            run.text = ""
    if first_run is None:
        paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        first_run = paragraph.add_run()
    first_run.text = "" if value is None else str(value)


def _find_docx_table(document, expected_headers, fallback_index=None):
    expected = tuple(_normalise_docx_header(v) for v in expected_headers)
    for table in document.tables:
        if not table.rows:
            continue
        actual = tuple(_normalise_docx_header(c.text) for c in table.rows[0].cells)
        if actual == expected:
            return table
    if fallback_index is not None and 0 <= fallback_index < len(document.tables):
        table = document.tables[fallback_index]
        if table.rows and len(table.rows[0].cells) == len(expected):
            return table
    raise ValidationError("Không tìm thấy bảng động đúng cấu trúc trong template.")


def _apply_docx_dynamic_tables(document, specs):
    """Clone one styled prototype row into exactly N deterministic rows.

    A spec contains only structural metadata and already-authorised row values.  This
    helper deliberately does not log values because personnel rows contain PII.
    """
    import copy

    for spec in specs or []:
        table = _find_docx_table(document, spec["headers"], spec.get("fallback_index"))
        if len(table.rows) < 2:
            raise ValidationError("Template thiếu dòng mẫu cho bảng động.")
        prototype = copy.deepcopy(table.rows[1]._tr)
        for row in list(table.rows[1:]):
            table._tbl.remove(row._tr)
        columns = tuple(spec["columns"])
        for values in spec.get("rows") or []:
            table._tbl.append(copy.deepcopy(prototype))
            cells = table.rows[-1].cells
            if len(cells) != len(columns):
                raise ValidationError("Số cột bảng động không khớp cấu hình template.")
            for cell, key in zip(cells, columns):
                _set_docx_cell_text_preserve_style(cell, values.get(key, ""))
            bold_row_key = spec.get("bold_row_key")
            if bold_row_key and values.get(bold_row_key):
                for cell in cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True


def _export_docx(tpl_path, m, items, dynamic_tables=None):
    import copy
    import docx
    if tpl_path and os.path.isfile(tpl_path):
        d = docx.Document(tpl_path)
        _apply_docx_dynamic_tables(d, dynamic_tables)
        # WO-29fix: NHO hang mau (XML) cua bang dong-danh-so TRUOC khi fill — de nhan ban
        # them dong khi so hang muc NHIEU HON so dong template (vd template BBNT 2 dong,
        # bao gia 4 hang muc -> truoc day 2 hang muc cuoi bi rot mat).
        extra_tables = []   # (tbl, so_dong_mau, tr_xml_mau)
        for tbl in d.tables:
            max_n, pattern_tr = 0, None
            for row in tbl.rows:
                ns = [int(mm.group(2)) for cell in row.cells
                      for mm in _IDX_TOKEN_RE.finditer(cell.text)]
                if ns and max(ns) > max_n:
                    max_n, pattern_tr = max(ns), copy.deepcopy(row._tr)
            if pattern_tr is not None and len(items) > max_n:
                extra_tables.append((tbl, max_n, pattern_tr))
        cursor = [0]
        for p in d.paragraphs:
            if "{{" in p.text:
                new = _fill_text(p.text, m, items, cursor)
                if new != p.text:
                    for r in p.runs:
                        r.text = ""
                    if p.runs:
                        p.runs[0].text = new
                    else:
                        p.add_run(new)
        for tbl in d.tables:
            for row in tbl.rows:
                row_has_item = any("{{" in cell.text and re.search(
                    r"quotation_item|item_name|model|item", cell.text) for cell in row.cells)
                for cell in row.cells:
                    if "{{" in cell.text:
                        new = _fill_text(cell.text, m, items, cursor)
                        cell.text = new
                if row_has_item:
                    cursor[0] += 1
        # them dong cho hang muc vuot qua so dong mau cua template
        for tbl, max_n, pattern_tr in extra_tables:
            for i in range(max_n, len(items)):
                tbl._tbl.append(copy.deepcopy(pattern_tr))
                for cell in tbl.rows[-1].cells:
                    txt = cell.text
                    if not txt.strip():
                        continue
                    new = _IDX_TOKEN_RE.sub(lambda mm: str(_item_field(items[i], mm.group(1))), txt)
                    if new == txt:                     # o tinh (khong co token)
                        st = txt.strip()
                        if st.isdigit():
                            new = str(i + 1)           # cot STT
                        elif "Đạt" in st and items[i].get("ket_qua"):
                            new = str(items[i]["ket_qua"])
                    if new != txt:
                        cell.text = new
    else:
        d = docx.Document()
        d.add_heading(m.get("company_name", ""), level=2)
        d.add_paragraph("MST: %s — %s" % (m.get("company_tax_id", ""), m.get("company_address", "")))
        d.add_heading(m.get("_title", "CHỨNG TỪ"), level=1)
        d.add_paragraph("Khách hàng: %s — Ngày: %s" % (m.get("customer_name", ""), m.get("date", "")))
        if items:
            t = d.add_table(rows=1, cols=4)
            hdr = t.rows[0].cells
            for i, h in enumerate(["Hạng mục", "Khối lượng", "Đơn giá", "Thành tiền"]):
                hdr[i].text = h
            for it in items:
                r = t.add_row().cells
                r[0].text = str(it.get("hang_muc") or it.get("ten_hang") or "")
                r[1].text = str(it.get("khoi_luong") or "")
                r[2].text = fmt_vnd(it.get("don_gia"))
                r[3].text = fmt_vnd(it.get("thanh_tien"))
        d.add_paragraph("")
        d.add_paragraph("ĐẠI DIỆN BÊN A\t\t\t\tĐẠI DIỆN BÊN B")
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


DOC_TABLES = {
    "quotation": ("quotation", "quotation_item", "quotation_id", "Báo giá"),
    "bbnt": ("bbnt", "bbnt_item", "bbnt_id", "Biên bản nghiệm thu"),
    "bqt": ("bqt", "bqt_item", "bqt_id", "Bảng quyết toán"),
    "payment": ("payment_request", None, None, "Thư đề nghị thanh toán"),
    "dccn": ("dccn", None, None, "Biên bản đối chiếu công nợ"),
    "pxk": ("pxk", "pxk_dong", "pxk_id", "Phiếu xuất kho"),
    "checklist": ("checklist_ct", "checklist_dong", "checklist_id", "Checklist nghiệm thu"),
    "hop_dong": ("hop_dong_ct", None, None, "Hợp đồng"),
}


# ==================== WO-34A: HO SO CONG TRINH (CT-00..CT-09) =============
# Nhanh SONG SONG voi DOC_TABLES/_data_map cu (khong dung chung, khong sua cu).
# Template bundle + mapping 28 ma CT-* doc tu TEMPLATE_MAPPING_CONG_TRINH_APP8777.json.
TPL_ROOT_CT = (r"D:\Quản trị DOANH NGHIỆP\Mẫu chứng từ chuẩn"
               r"\_MẪU HỒ SƠ CÔNG TRÌNH"
               r"\TH_ERP_V3_1")
_CT_INTEGRATION_DIR = "90_ERP_INTEGRATION"
_CT_MAPPING_FILE = "TEMPLATE_MAPPING_CONG_TRINH_APP8777_V3.json"
_CT_REQUIRED_RULES_FILE = "REQUIRED_DOCUMENT_RULES.json"
_CT_TPL_CACHE = None
_CT_REQUIRED_CACHE = None


def _resolve_ct_template_path(rel_path):
    """Tra path that; sua mem 11 path XLSX stale trong manifest V2.

    Mapping V2 con chen folder `BỘ 11 MẪU BÁO GIÁ RIÊNG` trong khi cac file da
    duoc dat phang o `Báo giá gửi chủ đầu tư`. Khong sua goi mau cua nguoi dung;
    resolver chi fallback theo basename duy nhat va bao `path_resolved` cho UI.
    """
    rel = (rel_path or "").replace("/", os.sep)
    direct = os.path.join(TPL_ROOT_CT, rel)
    if os.path.isfile(direct):
        return direct
    basename = os.path.basename(rel)
    hits = glob.glob(os.path.join(TPL_ROOT_CT, "**", basename), recursive=True)
    return hits[0] if len(hits) == 1 else direct


def _ct_templates_v2_legacy():
    """Registry V2: 29 mau goc + 27 mau bo sung, gom theo code."""
    global _CT_TPL_CACHE
    if _CT_TPL_CACHE is None:
        import json
        mp = os.path.join(TPL_ROOT_CT, "00. HƯỚNG DẪN & CHECKLIST",
                          "TEMPLATE_MAPPING_CONG_TRINH_APP8777_V2.json")
        try:
            with open(mp, encoding="utf-8") as f:
                data = json.load(f)
            entries = list(data.get("templates", [])) + list(data.get("additional_templates_v2", []))
            out = {}
            for t in entries:
                path = _resolve_ct_template_path(t.get("path"))
                out[t["code"]] = {
                    "title": t.get("title") or t["code"], "path": t.get("path") or "",
                    "abs_path": path, "exists": os.path.isfile(path),
                    "format": os.path.splitext(path)[1].lower().lstrip(".") or "docx",
                    "source": t.get("source") or "V2",
                }
            _CT_TPL_CACHE = out
        except OSError:
            _CT_TPL_CACHE = {}
    return _CT_TPL_CACHE


def ct_templates():
    """Return the authoritative V3.1 registry, keyed by document code.

    New V3.1 templates become visible from the data registry immediately, but
    rendering remains explicitly gated to structured generators below.  That
    prevents a template with an unwired dynamic table from being exported blank.
    """
    global _CT_TPL_CACHE
    if _CT_TPL_CACHE is None:
        import json
        mapping_path = os.path.join(TPL_ROOT_CT, _CT_INTEGRATION_DIR, _CT_MAPPING_FILE)
        try:
            with open(mapping_path, encoding="utf-8") as handle:
                mapping = json.load(handle)
            quote_contract_by_file = {}
            contract_rel = mapping.get("quotation_dynamic_contract")
            if contract_rel:
                contract_path = os.path.join(TPL_ROOT_CT, contract_rel.replace("/", os.sep))
                with open(contract_path, encoding="utf-8") as handle:
                    contract = json.load(handle)
                quote_contract_by_file = {
                    str(row.get("file") or ""): dict(row)
                    for row in contract.get("templates", [])
                }
            registry = {}
            for entry in mapping.get("templates", []):
                path = _resolve_ct_template_path(entry.get("path"))
                registry[entry["code"]] = {
                    "title": entry.get("title") or entry["code"],
                    "path": entry.get("path") or "",
                    "abs_path": path,
                    "exists": os.path.isfile(path),
                    "format": (entry.get("file_type") or
                               os.path.splitext(path)[1].lower().lstrip(".") or "docx").lower(),
                    "source": "V3.1",
                    "registry_version": str(mapping.get("version") or "3.1"),
                    "template_id": entry.get("template_id"),
                    "legacy_code": entry.get("legacy_code"),
                    "phase_code": entry.get("phase_code"),
                    "owner_role": entry.get("owner_role"),
                    "reviewer_role": entry.get("reviewer_role"),
                    "approver_role": entry.get("approver_role"),
                    "output_formats": list(entry.get("output_formats") or []),
                    "dynamic_rows": quote_contract_by_file.get(os.path.basename(path)),
                }
            _CT_TPL_CACHE = registry
        except (OSError, ValueError, KeyError):
            _CT_TPL_CACHE = {}
    return _CT_TPL_CACHE


def ct_document_requirements(profile_code="INSTALLATION_STANDARD"):
    """Return V3.1 required/conditional codes for a project template profile.

    Conditional documents are exposed to users but are not counted as missing
    until their business trigger is enabled for the project.  This is important:
    V3.1 has 84 templates, not 84 mandatory documents for every project.
    """
    global _CT_REQUIRED_CACHE
    if _CT_REQUIRED_CACHE is None:
        import json
        rules_path = os.path.join(TPL_ROOT_CT, _CT_INTEGRATION_DIR,
                                  _CT_REQUIRED_RULES_FILE)
        try:
            with open(rules_path, encoding="utf-8") as handle:
                _CT_REQUIRED_CACHE = json.load(handle)
        except (OSError, ValueError):
            _CT_REQUIRED_CACHE = {}
    profiles = _CT_REQUIRED_CACHE.get("project_profiles") or {}
    selected = profiles.get(profile_code) or profiles.get("INSTALLATION_STANDARD") or {}
    return {
        "profile_code": profile_code if profile_code in profiles else "INSTALLATION_STANDARD",
        "required": tuple(selected.get("required") or ()),
        "conditional": tuple(selected.get("conditional") or ()),
        "conditional_triggers": tuple(_CT_REQUIRED_CACHE.get("conditional_triggers") or ()),
    }


def ct_document_profiles():
    """Return the V3.1 project-profile codes available to the importer UI."""
    # Populate the same cached rules source without duplicating its file logic.
    ct_document_requirements()
    return tuple(sorted((_CT_REQUIRED_CACHE.get("project_profiles") or {}).keys()))


_CT_PERSONNEL_TABLES = {
    "CT-01-DSNS": {
        "headers": ("STT", "Họ và tên", "Chức vụ/Nhiệm vụ", "CCCD/CMND",
                    "Năm sinh", "Số điện thoại", "Ghi chú"),
        "columns": ("stt", "ho_ten", "chuc_vu", "cccd", "nam_sinh", "sdt", "ghi_chu"),
        "fallback_index": 2,
    },
    "CT-01-PKBNV": {
        "headers": ("STT", "Họ và tên", "Ngày/tháng/năm sinh", "CCCD",
                    "Số điện thoại", "Chức vụ", "Ghi chú"),
        "columns": ("stt", "ho_ten", "ngay_sinh", "cccd", "sdt", "chuc_vu", "ghi_chu"),
        "fallback_index": 2,
    },
}


_CT_BOQ_APPENDIX_TABLES = {
    # V3.1 canonical code.  The old code remains below for historical draft
    # records only; both use the same exact imported BOQ hierarchy.
    "HD-07": {
        "headers": ("STT", "TÃªn hÃ ng hÃ³a / háº¡ng má»¥c", "Quy cÃ¡ch / Model", "ÄVT",
                    "Sá»‘ lÆ°á»£ng", "ÄÆ¡n giÃ¡", "ThÃ nh tiá»n", "Ghi chÃº"),
        "columns": ("stt", "ten_hang_muc", "quy_cach", "dvt", "so_luong",
                    "don_gia", "thanh_tien", "ghi_chu"),
        "fallback_index": 1,
        "bold_row_key": "is_heading",
    },
    # Actual V2 HD-07 appendix: no VAT column.  VAT remains in the imported BOQ
    # record and is not silently folded into the note or pre-tax amount column.
    "HD-07_PHU_LUC_BANG_KHOI_LUO": {
        "headers": ("STT", "Tên hàng hóa / hạng mục", "Quy cách / Model", "ĐVT",
                    "Số lượng", "Đơn giá", "Thành tiền", "Ghi chú"),
        "columns": ("stt", "ten_hang_muc", "quy_cach", "dvt", "so_luong",
                    "don_gia", "thanh_tien", "ghi_chu"),
        "fallback_index": 1,
        "bold_row_key": "is_heading",
    },
}


# Batch 6: only the two non-financial V3.1 acceptance minutes are wired here.
# CT-08-QTHT contains prices and remains fail-closed until the finance/signing batch.
_CT_ACCEPTANCE_TABLES = {
    "CT-06-BBNTGD": {
        "main": {
            "headers": ("STT", "Giai đoạn/Bộ phận", "Từ ngày", "Đến ngày",
                        "Khối lượng nghiệm thu", "Kết quả"),
            "columns": ("stt", "stage_item", "period_from", "period_to",
                        "acceptance_display", "result"),
            "fallback_index": 2,
        },
        "issues": {
            "headers": ("STT", "Tồn tại", "Biện pháp", "Người phụ trách", "Hạn xử lý"),
            "columns": ("stt", "issue", "measure", "owner", "deadline"),
            "fallback_index": 3,
        },
    },
    "CT-06-BBNTHH": {
        "main": {
            "headers": ("STT", "Hạng mục", "Giá trị/KL theo hợp đồng", "KL hoàn thành",
                        "Tài liệu nghiệm thu kèm theo", "Kết luận"),
            "columns": ("stt", "stage_item", "planned_display", "acceptance_display",
                        "evidence", "conclusion"),
            "fallback_index": 2,
        },
        "issues": {
            "headers": ("STT", "Nội dung tồn tại", "Ảnh/chứng cứ", "Biện pháp xử lý",
                        "Thời hạn"),
            "columns": ("stt", "issue", "evidence", "measure", "deadline"),
            "fallback_index": 3,
        },
    },
}


# CT-05-NKTC is a repeating document: one approved journal record produces one
# versioned artifact.  The dossier layer verifies that every approved record is
# covered before the project may pass acceptance readiness.
_CT_JOURNAL_TABLES = {
    "CT-05-NKTC": {
        "daily": {
            "headers": ("Thời tiết", "Nhân lực", "Thiết bị", "Vật tư sử dụng",
                        "Thời gian làm việc"),
            "columns": ("weather", "workforce", "equipment", "materials", "work_hours"),
            "fallback_index": 2,
        },
        "work": {
            "headers": ("STT", "Khu vực/Tầng", "Nội dung công việc", "Khối lượng",
                        "Người thực hiện", "Kết quả"),
            "columns": ("stt", "stage", "content", "quantity", "worker", "result"),
            "fallback_index": 3,
        },
        "issues": {
            "headers": ("STT", "Nội dung", "Ảnh minh chứng", "Biện pháp xử lý",
                        "Người phụ trách", "Hạn xử lý"),
            "columns": ("stt", "content", "evidence", "measure", "owner", "deadline"),
            "fallback_index": 4,
        },
    }
}


def _sqlite_table_exists(conn, table):
    return bool(conn.execute("SELECT 1 FROM sqlite_master WHERE type='table' AND name=?",
                             (table,)).fetchone())


def _sqlite_columns(conn, table):
    if not _sqlite_table_exists(conn, table):
        return set()
    return {row[1] for row in conn.execute("PRAGMA table_info(%s)" % table).fetchall()}


def _row_value(row, *keys):
    if row is None:
        return ""
    available = set(row.keys()) if hasattr(row, "keys") else set()
    for key in keys:
        if available and key not in available:
            continue
        try:
            value = row[key]
        except (KeyError, IndexError, TypeError):
            continue
        if value not in (None, ""):
            return value
    return ""


def _ct_person_date(value, year_only=False):
    text = str(value or "").strip()
    if not text:
        return ""
    years = re.findall(r"(?<!\d)(?:19|20)\d{2}(?!\d)", text)
    if year_only:
        return years[-1] if years else text
    if re.fullmatch(r"\d{4}-\d{1,2}-\d{1,2}", text):
        return fmt_d(text)
    return text


def _ct_project_personnel(conn, project_id):
    """Authorised project assignments joined to the personnel master record.

    Never falls back to exporting every employee: an absent assignment must not
    broaden the PII scope of a project document.
    """
    if not (_sqlite_table_exists(conn, "project_personnel") and
            _sqlite_table_exists(conn, "nhan_su")):
        return []
    rows = conn.execute("""SELECT pp.*, n.ho_ten AS _ho_ten, n.loai AS _loai,
                                  n.sdt AS _sdt, n.cccd AS _cccd,
                                  n.ngay_sinh AS _ngay_sinh
                           FROM project_personnel pp
                           JOIN nhan_su n ON n.id=pp.nhan_su_id
                           WHERE pp.project_id=?
                           ORDER BY CASE WHEN pp.source_row IS NULL THEN 1 ELSE 0 END,
                                    pp.source_row, pp.id""", (project_id,)).fetchall()
    result = []
    for index, row in enumerate(rows, 1):
        role = _row_value(row, "project_role", "site_role", "_loai")
        result.append({
            "stt": _row_value(row, "source_stt") or index,
            "ho_ten": _row_value(row, "_ho_ten"),
            "chuc_vu": role,
            "cccd": str(_row_value(row, "_cccd")),
            "nam_sinh": _ct_person_date(_row_value(row, "_ngay_sinh"), year_only=True),
            "ngay_sinh": _ct_person_date(_row_value(row, "_ngay_sinh")),
            "sdt": str(_row_value(row, "_sdt")),
            "ghi_chu": str(_row_value(row, "source_note")),
        })
    return result


def _ct_personnel_dynamic_spec(ma_mau, rows):
    config = _CT_PERSONNEL_TABLES.get(ma_mau)
    if not config:
        return []
    spec = dict(config)
    spec["rows"] = list(rows or [])
    return [spec]


def _ct_boq_display_number(raw_value, numeric_value, *, money=False):
    """Render a BOQ number without ever exposing an Excel formula as document text."""
    raw = str(raw_value or "").strip()
    # A literal source value is more faithful than a REAL round-trip.  Formula
    # text is never exposed; for formulas use the parser's canonical result.
    if raw and not raw.startswith("="):
        number = _ct_number(raw)
        if number is not None:
            return fmt_vnd(number) if money else _fmt_qty(raw)
        return raw
    if numeric_value not in (None, ""):
        return fmt_vnd(numeric_value) if money else _fmt_qty(numeric_value)
    if not raw or raw.startswith("="):
        return ""
    number = _ct_number(raw)
    if number is not None:
        return fmt_vnd(number) if money else _fmt_qty(number)
    return raw


def _ct_profile_boq_rows(conn, project_id):
    """Return the exact active imported BOQ hierarchy for the HD-07 appendix.

    This deliberately has no legacy quotation fallback: HD-07 is a contractual
    appendix and must never be generated from the old predicted-material rows.
    """
    if not (_sqlite_table_exists(conn, "project_profile_import") and
            _sqlite_table_exists(conn, "project_boq_line")):
        return None, []
    profile = conn.execute("""SELECT * FROM project_profile_import
                              WHERE project_id=? AND status='active'
                              ORDER BY id DESC LIMIT 1""", (project_id,)).fetchone()
    if not profile:
        return None, []
    lines = conn.execute("""SELECT * FROM project_boq_line
                             WHERE profile_import_id=?
                             ORDER BY thu_tu, source_row, id""", (profile["id"],)).fetchall()
    rows = []
    for line in lines:
        is_heading = str(_row_value(line, "line_type")).strip().lower() == "heading"
        rows.append({
            "stt": str(_row_value(line, "source_stt_raw")),
            "ten_hang_muc": str(_row_value(line, "item_name_raw")),
            "quy_cach": "" if is_heading else str(_row_value(line, "technical_requirement_raw")),
            "dvt": "" if is_heading else str(_row_value(line, "uom_raw")),
            "so_luong": "" if is_heading else _ct_boq_display_number(
                _row_value(line, "contract_qty_raw"), _row_value(line, "contract_qty")),
            "don_gia": "" if is_heading else _ct_boq_display_number(
                _row_value(line, "unit_price_raw"), _row_value(line, "unit_price"), money=True),
            "thanh_tien": "" if is_heading else _ct_boq_display_number(
                _row_value(line, "amount_raw"), _row_value(line, "amount"), money=True),
            "ghi_chu": "" if is_heading else str(_row_value(line, "note_raw")),
            # Kept for exact source fidelity/future 9-column template revisions.
            "vat_rate": "" if is_heading else str(_row_value(line, "vat_rate_raw")),
            "is_heading": is_heading,
            "hierarchy_level": _row_value(line, "hierarchy_level"),
        })
    return profile, rows


def _ct_boq_dynamic_spec(ma_mau, rows):
    config = _CT_BOQ_APPENDIX_TABLES.get(ma_mau)
    if not config:
        return []
    spec = dict(config)
    spec["rows"] = list(rows or [])
    return [spec]


def _ct_acceptance_dynamic_spec(ma_mau, rows, issues):
    """Return exact table contracts for the two Batch 6 V3.1 BBNT templates."""
    config = _CT_ACCEPTANCE_TABLES.get(ma_mau)
    if not config:
        return []
    main = dict(config["main"])
    main["rows"] = list(rows or [])
    issue_spec = dict(config["issues"])
    issue_spec["rows"] = list(issues or [])
    return [main, issue_spec]


def _ct_journal_doc_data(conn, project_id, journal_id):
    """Return scalar/dynamic data for one approved daily journal.

    Fail closed when a required template field was not captured.  Exporting a
    pretty but incomplete journal would be more dangerous than showing a clear
    validation error before acceptance.
    """
    try:
        journal_id = int(journal_id)
    except (TypeError, ValueError):
        raise ValidationError("Phải chọn một nhật ký hợp lệ để xuất CT-05-NKTC.")
    row = conn.execute("""SELECT n.*,u.full_name AS creator_name,
            q.planned_qty,l.item_name_raw,l.uom_raw,s.name_raw AS stage_name
        FROM nhat_ky_thi_cong n
        LEFT JOIN app_user u ON u.id=n.created_by
        LEFT JOIN project_boq_stage_qty q ON q.id=n.boq_stage_qty_id
        LEFT JOIN project_boq_line l ON l.id=q.boq_line_id
        LEFT JOIN project_boq_stage s ON s.id=q.stage_id
        WHERE n.id=? AND n.project_id=?""", (journal_id, project_id)).fetchone()
    if not row:
        raise ValidationError("Nhật ký không thuộc đúng công trình.")
    if _row_value(row, "trang_thai") != "Da_duyet":
        raise ValidationError("Chỉ xuất bản trình chủ đầu tư từ nhật ký đã được KTT duyệt.")
    required = {
        "nhan_luc": _row_value(row, "nhan_luc"),
        "thoi_gian_lam_viec": _row_value(row, "thoi_gian_lam_viec"),
        "ket_qua": _row_value(row, "ket_qua"),
        "noi_dung": _row_value(row, "noi_dung"),
        "khoi_luong_thuc_hien": _row_value(row, "khoi_luong_thuc_hien"),
        "boq_stage_qty_id_or_manual_item": (_row_value(row, "boq_stage_qty_id")
                                              or _row_value(row, "hang_muc_tu_do")),
    }
    missing = [key for key, value in required.items() if value in (None, "")]
    equipment = str(_row_value(row, "thiet_bi")).strip()
    if not equipment and not bool(_row_value(row, "khong_su_dung_thiet_bi")):
        missing.append("thiet_bi")
    if missing:
        raise ValidationError("Nhật ký chưa đủ dữ liệu để xuất mẫu chuẩn: %s." % ", ".join(missing))

    materials = conn.execute("""SELECT ten_vat_tu,dvt,so_luong_su_dung,ghi_chu
        FROM nhat_ky_vat_tu WHERE nhat_ky_id=? ORDER BY id""", (journal_id,)).fetchall()
    if materials:
        material_text = "; ".join(
            "%s: %s%s%s" % (
                _row_value(item, "ten_vat_tu"),
                _fmt_qty(_row_value(item, "so_luong_su_dung")),
                (" " + str(_row_value(item, "dvt"))) if _row_value(item, "dvt") else "",
                (" (" + str(_row_value(item, "ghi_chu")) + ")") if _row_value(item, "ghi_chu") else "",
            ) for item in materials)
    elif bool(_row_value(row, "khong_su_dung_vat_tu")):
        material_text = "Không sử dụng vật tư"
    else:
        raise ValidationError("Nhật ký chưa có vật tư sử dụng hoặc xác nhận không sử dụng vật tư.")

    photos = conn.execute("""SELECT h.file_anh,sd.file_name
        FROM cong_trinh_hinh_anh h
        LEFT JOIN source_document sd ON sd.id=h.source_document_id
        WHERE h.nhat_ky_id=? ORDER BY h.id""", (journal_id,)).fetchall()
    photo_names = [os.path.basename(str(_row_value(item, "file_name", "file_anh")))
                   for item in photos if _row_value(item, "file_name", "file_anh")]
    evidence = ", ".join(photo_names)
    stage = " — ".join(part for part in (
        str(_row_value(row, "stage_name")).strip(),
        str(_row_value(row, "item_name_raw")).strip()) if part)
    if not stage:
        stage = "Nhật ký tổng quát - " + str(_row_value(row, "hang_muc_tu_do")).strip()
    uom = str(_row_value(row, "uom_raw")).strip()
    quantity = (_fmt_qty(_row_value(row, "khoi_luong_thuc_hien")) +
                ((" " + uom) if uom else ""))
    daily = [{
        "weather": str(_row_value(row, "thoi_tiet")),
        "workforce": str(_row_value(row, "nhan_luc")),
        "equipment": equipment or "Không sử dụng thiết bị",
        "materials": material_text,
        "work_hours": str(_row_value(row, "thoi_gian_lam_viec")),
    }]
    work = [{
        "stt": 1, "stage": stage, "content": str(_row_value(row, "noi_dung")),
        "quantity": quantity, "worker": str(_row_value(row, "creator_name")),
        "result": str(_row_value(row, "ket_qua")) +
                  (("; Ảnh: " + evidence) if evidence else ""),
    }]
    issue_text = "; ".join(part for part in (
        str(_row_value(row, "su_co")).strip(),
        str(_row_value(row, "kho_khan_kien_nghi")).strip()
        if not bool(_row_value(row, "khong_co_kien_nghi")) else "") if part)
    issues = []
    if issue_text:
        issue_required = {
            "bien_phap_xu_ly": _row_value(row, "bien_phap_xu_ly"),
            "nguoi_phu_trach_xu_ly": _row_value(row, "nguoi_phu_trach_xu_ly"),
            "han_xu_ly": _row_value(row, "han_xu_ly"),
        }
        issue_missing = [key for key, value in issue_required.items() if value in (None, "")]
        if issue_missing:
            raise ValidationError("Nhật ký có vướng mắc nhưng thiếu dữ liệu xử lý: %s." %
                                  ", ".join(issue_missing))
        issues.append({
            "stt": 1, "content": issue_text, "evidence": evidence,
            "measure": str(_row_value(row, "bien_phap_xu_ly")),
            "owner": str(_row_value(row, "nguoi_phu_trach_xu_ly")),
            "deadline": fmt_d(_row_value(row, "han_xu_ly")),
        })
    config = _CT_JOURNAL_TABLES["CT-05-NKTC"]
    specs = []
    for name, values in (("daily", daily), ("work", work), ("issues", issues)):
        spec = dict(config[name])
        spec["rows"] = values
        specs.append(spec)
    scalar = {
        "HANG_MUC_THI_CONG": stage,
        "NGAY_LAP": fmt_d(_row_value(row, "ngay_ghi")),
        "NGUOI_LAP": str(_row_value(row, "creator_name")),
        "HO_TEN": str(_row_value(row, "creator_name")),
    }
    return row, scalar, specs


def _ct_acceptance_doc_rows(conn, project_id, acceptance_id):
    acceptance = conn.execute("""SELECT * FROM project_acceptance
        WHERE id=? AND project_id=?""", (acceptance_id, project_id)).fetchone()
    if not acceptance:
        raise ValidationError("Đợt nghiệm thu không thuộc đúng công trình.")
    if acceptance["status"] not in ("Draft", "Can_bo_sung"):
        raise ValidationError("Chỉ sinh lại biên bản khi đợt nghiệm thu đang là bản nháp.")
    raw_rows = conn.execute("""SELECT ai.*,q.planned_qty,l.item_name_raw,l.uom_raw,
            s.name_raw AS stage_name
        FROM project_acceptance_item ai
        JOIN project_boq_stage_qty q ON q.id=ai.boq_stage_qty_id
        JOIN project_boq_line l ON l.id=q.boq_line_id
        JOIN project_boq_stage s ON s.id=q.stage_id
        WHERE ai.acceptance_id=? ORDER BY s.thu_tu,l.thu_tu,l.source_row,ai.id""",
                            (acceptance_id,)).fetchall()
    if not raw_rows:
        raise ValidationError("Đợt nghiệm thu chưa có dòng khối lượng exact BOQ.")
    rows, issues = [], []
    period_from = fmt_d(_row_value(acceptance, "period_from"))
    period_to = fmt_d(_row_value(acceptance, "period_to"))
    for index, row in enumerate(raw_rows, 1):
        uom = str(_row_value(row, "uom_raw"))
        accepted = _fmt_qty(_row_value(row, "acceptance_qty"))
        planned = _fmt_qty(_row_value(row, "planned_qty"))
        label = "%s — %s" % (_row_value(row, "stage_name"), _row_value(row, "item_name_raw"))
        rows.append({"stt": index, "stage_item": label,
                     "period_from": period_from, "period_to": period_to,
                     "planned_display": (planned + (" " + uom if uom else "")).strip(),
                     "acceptance_display": (accepted + (" " + uom if uom else "")).strip(),
                     "evidence": "Nhật ký đã duyệt / hồ sơ V3.1",
                     "result": "Đạt", "conclusion": "Đạt"})
        if abs(float(_row_value(row, "acceptance_qty") or 0) -
               float(_row_value(row, "journal_confirmed_qty") or 0)) > 1e-9:
            issues.append({"stt": len(issues) + 1,
                           "issue": str(_row_value(row, "discrepancy_reason")),
                           "evidence": "Đối chiếu khối lượng nhật ký",
                           "measure": "Đã xác nhận chênh lệch có lý do",
                           "owner": "KTT", "deadline": "Trước khi phát hành"})
    return acceptance, rows, issues


def _ct_latest_quotation(conn, project_id):
    columns = _sqlite_columns(conn, "quotation")
    if not columns:
        return None
    where = ["project_id=?"]
    if "status" in columns:
        where.append("COALESCE(status,'')<>'Huy'")
    order = []
    if "is_official" in columns:
        order.append("COALESCE(is_official,0) DESC")
    if "imported_at" in columns:
        order.append("CASE WHEN imported_at IS NULL THEN 1 ELSE 0 END")
        order.append("imported_at DESC")
    order.append("id DESC")
    return conn.execute("SELECT * FROM quotation WHERE %s ORDER BY %s LIMIT 1" %
                        (" AND ".join(where), ", ".join(order)), (project_id,)).fetchone()


def _ct_quotation_items(conn, project_id):
    quotation = _ct_latest_quotation(conn, project_id)
    if not quotation:
        return None, []
    columns = _sqlite_columns(conn, "quotation_item")
    order = []
    if "source_row" in columns:
        order += ["CASE WHEN source_row IS NULL THEN 1 ELSE 0 END", "source_row"]
    if "stt" in columns:
        order.append("stt")
    order.append("id")
    rows = [dict(row) for row in conn.execute(
        "SELECT * FROM quotation_item WHERE quotation_id=? ORDER BY %s" % ", ".join(order),
        (_row_value(quotation, "id"),)).fetchall()]

    # Enrich with exact source fields when the profile importer has linked a BOQ line.
    if (_sqlite_table_exists(conn, "project_profile_import") and
            _sqlite_table_exists(conn, "project_boq_line")):
        profile = conn.execute("""SELECT id FROM project_profile_import
                                  WHERE project_id=? AND quotation_id=? AND status='active'
                                  ORDER BY id DESC LIMIT 1""",
                               (project_id, _row_value(quotation, "id"))).fetchone()
        if profile:
            linked = {}
            for line in conn.execute("""SELECT * FROM project_boq_line
                                        WHERE profile_import_id=? AND quotation_item_id IS NOT NULL""",
                                     (_row_value(profile, "id"),)).fetchall():
                linked[_row_value(line, "quotation_item_id")] = dict(line)
            for item in rows:
                line = linked.get(item.get("id"))
                if not line:
                    continue
                aliases = {
                    "source_stt": "source_stt_raw",
                    "source_item_raw": "item_name_raw",
                    "technical_requirement": "technical_requirement_raw",
                    "dvt": "uom_raw",
                    "quantity_raw": "contract_qty_raw",
                    "unit_price_raw": "unit_price_raw",
                    "brand_raw": "brand_raw",
                    "vat_rate_raw": "vat_rate_raw",
                    "source_note_raw": "note_raw",
                    "line_type": "line_type",
                }
                for target, source in aliases.items():
                    if item.get(target) in (None, "") and line.get(source) not in (None, ""):
                        item[target] = line[source]
    return quotation, rows


def _ct_data_map(conn, project, sess, extra=None, quotation=None, contract=None):
    """Map {{TOKEN_HOA}} cho nhom cong trinh. Nguon: project + customer + cau_hinh
    + hop_dong (qua quotation cua project) + tien_do. Token la CHU HOA theo dung
    placeholder_policy cua bundle ({{TEN_CONG_TRINH}}...)."""
    kh = conn.execute("SELECT * FROM customer WHERE id=?", (project["customer_id"],)).fetchone()
    cfg_row = conn.execute("SELECT * FROM cau_hinh WHERE id=1").fetchone()
    cfg = dict(cfg_row) if cfg_row else {}
    hd = contract
    if hd is None:
        hd = conn.execute("""SELECT h.* FROM hop_dong_ct h
                             JOIN quotation q ON q.id=h.quotation_id
                             WHERE q.project_id=? ORDER BY h.id DESC LIMIT 1""",
                          (project["id"],)).fetchone()
    td = conn.execute("""SELECT hang_muc FROM cong_trinh_tien_do WHERE project_id=?
                         ORDER BY id LIMIT 1""", (project["id"],)).fetchone()
    quotation = quotation or _ct_latest_quotation(conn, project["id"])
    today = date.today()
    project_name = _row_value(project, "project_name")
    project_location = _row_value(project, "dia_diem", "khu_vuc")
    customer_name = _row_value(kh, "customer_name")
    customer_address = _row_value(kh, "dia_chi")
    customer_tax_id = _row_value(kh, "tax_id")
    contact_person = _row_value(kh, "nguoi_lien_he")
    contract_no = _row_value(hd, "code")
    contract_date = fmt_d(_row_value(hd, "ngay_ky"))
    work_item = _row_value(td, "hang_muc") or project_name
    quotation_no = _row_value(quotation, "code")
    quotation_date = fmt_d(_row_value(quotation, "ngay_lap"))
    quotation_total = _row_value(quotation, "grand_total")
    m = {
        "TEN_CONG_TRINH": project_name,
        "DIA_DIEM_CONG_TRINH": project_location,
        "TEN_CHU_DAU_TU": customer_name,
        "TEN_NHA_THAU": cfg.get("ten_cong_ty") or "",
        "MST_NHA_THAU": cfg.get("ma_so_thue") or "",
        "DIA_CHI_NHA_THAU": cfg.get("dia_chi") or "",
        "SO_HOP_DONG": contract_no,
        "NGAY_HOP_DONG": contract_date,
        "HANG_MUC_THI_CONG": work_item,
        "NGAY_LAP": fmt_d(today.isoformat()),
        "NGUOI_LAP": sess.get("full_name") or "",
        "HO_TEN": sess.get("full_name") or "",
        "LY_DO_KHAC": "",
        "MA_CONG_TRINH": _row_value(project, "code"),
        "NGAY": fmt_d(today.isoformat()),
        # Alias that are present in the V2 DOCX bundle but absent from its CSV dictionary.
        "CT": project_name, "DD": project_location, "CDT": customer_name,
        "HM": work_item, "HANG_MUC": work_item, "GOI": project_name,
        "SO_CAN_CU": contract_no, "HD": contract_no,
        "NOI": project_location, "NAM": str(today.year),
        "BQL": "", "TVGS": "", "TK": "",
        "BQL_REP": "", "TVGS_REP": "", "LAB_REP": "",
        # GIA_TRI_DE_NGHI: mặc định = grand_total BG leaf (CT-08-TDNTT).
        "GIA_TRI_DE_NGHI": (
            fmt_vnd(quotation_total) if quotation_total not in (None, "") else ""),
        "MO_TA_PHAM_VI_CONG_VIEC": project_name or "",
        "DUYET_TK": "",
        "SO_BAO_GIA": quotation_no or "",
        "NGAY_BAO_GIA": quotation_date or "",
        "GIA_TRI_BAO_GIA": (
            fmt_vnd(quotation_total) if quotation_total not in (None, "") else ""),
        # Lower-case scalar contract used by all 11 dynamic quotation workbooks.
        "customer_name": customer_name, "customer_address": customer_address,
        "customer_tax_id": customer_tax_id, "contact_person": contact_person,
        "project_name": project_name, "project_location": project_location,
        "quotation_id": quotation_no, "quote_no": quotation_no,
        "quote_date": quotation_date, "reference_no": contract_no,
        "valid_days": "30", "payment_terms": "Theo thỏa thuận hợp đồng",
        "delivery_terms": "", "warranty_terms": "",
        "amount_in_words": so_thanh_chu(quotation_total) if quotation_total not in (None, "") else "",
        # Contract-package aliases used by HD-07.  Unknown banking/signatory
        # fields stay blank unless the caller supplies them in ``extra``.
        "ben_a_ten": customer_name, "ben_a_dia_chi": customer_address,
        "ben_a_mst": customer_tax_id,
        "ben_a_dien_thoai": _row_value(kh, "dien_thoai"),
        "ben_a_email": _row_value(kh, "email"),
        "ben_a_dai_dien": contact_person, "ben_a_nguoi_ky": contact_person,
        "ben_a_chuc_vu": "", "ben_a_tai_khoan": "", "ben_a_ngan_hang": "",
        "ben_b_email": "", "ben_b_nguoi_ky": "",
        "so_hop_dong_goc": contract_no, "so_hop_dong": contract_no, "so_phu_luc": "",
        "dia_diem_ky": project_location,
        "dia_diem_thi_cong": project_location,
        "noi_dung_cong_viec": work_item or project_name,
        "gia_tri_hop_dong": (
            fmt_vnd(quotation_total) if quotation_total not in (None, "") else ""),
        "tong_thanh_toan": (
            fmt_vnd(quotation_total) if quotation_total not in (None, "") else ""),
        "ngay_hieu_luc": contract_date or fmt_d(today.isoformat()),
        "ngay_bat_dau": "",
        "ngay_hoan_thanh": "",
        "ngay_het_hieu_luc": "",
        "ngay": str(today.day), "thang": str(today.month), "nam": str(today.year),
    }
    for k, v in (extra or {}).items():
        m[str(k)] = "" if v is None else str(v)
    return m


def _validate_hd07_payment_schedule(mapping):
    """Reject a legal appendix whose payment schedule would be silently blank."""
    core = ("giai_doan", "ty_le_dot", "so_tien_dot", "dieu_kien", "ho_so_dot")
    populated = 0
    for index in range(1, 5):
        values = {name: str(mapping.get("%s_%d" % (name, index), "") or "").strip()
                  for name in core}
        if not any(values.values()):
            continue
        populated += 1
        missing = [name for name, value in values.items() if not value]
        if missing:
            raise ValidationError(
                "Lịch thanh toán đợt %d còn thiếu: %s." % (index, ", ".join(missing)))
    if not populated:
        raise ValidationError(
            "HD-07 chưa có lịch thanh toán có cấu trúc; không sinh phụ lục rỗng.")


def _ct_number(value):
    """Parse DB numerics and conservative Vietnamese-formatted numeric text."""
    if value in (None, "") or isinstance(value, bool):
        return None
    if isinstance(value, (int, float)):
        return float(value)
    text = str(value).strip().replace("\u00a0", "").replace(" ", "")
    text = re.sub(r"[^0-9,\.\-+]", "", text)
    if not text or text in ("-", "+"):
        return None
    if "," in text and "." in text:
        if text.rfind(",") > text.rfind("."):
            text = text.replace(".", "").replace(",", ".")
        else:
            text = text.replace(",", "")
    elif "," in text:
        parts = text.split(",")
        if len(parts) > 2 and all(len(part) == 3 for part in parts[1:]):
            text = "".join(parts)
        else:
            text = "".join(parts[:-1]) + "." + parts[-1] if len(parts[-1]) <= 3 else "".join(parts)
    elif text.count(".") > 1:
        parts = text.split(".")
        if all(len(part) == 3 for part in parts[1:]):
            text = "".join(parts)
        else:
            text = "".join(parts[:-1]) + "." + parts[-1]
    elif "." in text and len(text.rsplit(".", 1)[1]) == 3:
        text = text.replace(".", "")
    try:
        return float(text)
    except ValueError:
        return None


def _ct_vat_fraction(value):
    if value in (None, ""):
        return None
    text = str(value).strip()
    number = _ct_number(text.rstrip("%"))
    if number is None:
        return None
    if text.endswith("%") or abs(number) > 1:
        number /= 100.0
    return number


def _ct_qty_uom(item):
    # Linked BOQ source fields win over legacy quotation defaults (often numeric 0).
    uom = str(_row_value(item, "uom_raw", "dvt", "uom") or "").strip()
    raw = _row_value(item, "contract_qty_raw", "quantity_raw", "so_luong", "quantity",
                     "khoi_luong")
    if isinstance(raw, (int, float)) and not isinstance(raw, bool):
        return float(raw), uom
    text = str(raw or "").strip()
    match = re.match(r"^\s*([-+]?\d[\d.,\s]*)(?:\s+(.+?))?\s*$", text)
    if not match:
        return None, uom
    quantity = _ct_number(match.group(1))
    if not uom and match.group(2):
        uom = match.group(2).strip()
    return quantity, uom


def _ct_quote_row(item, index):
    quantity, uom = _ct_qty_uom(item)
    is_heading = str(_row_value(item, "line_type")).strip().lower() == "heading"
    # Source raw wins when available: legacy quotation rows may contain a default 0
    # even though the official BOQ line retained the actual unit price.
    price = _ct_number(_row_value(item, "unit_price_raw", "don_gia", "unit_price"))
    vat = _ct_vat_fraction(_row_value(item, "vat_rate_raw", "thue_suat", "vat_rate"))
    if is_heading:
        quantity = price = vat = None
    return [
        _row_value(item, "source_stt", "stt") or index,
        _row_value(item, "item_code_or_group", "ma_hang", "item_code", "nhom_hang"),
        _row_value(item, "source_item_raw", "item_name_raw", "hang_muc", "ten_hang"),
        _row_value(item, "technical_requirement", "technical_requirement_raw",
                   "quy_cach_model", "model", "spec", "brand_raw"),
        uom,
        quantity,
        price,
        vat,
    ]


def _shift_merged_rows(ws, start_row, delta):
    if not delta:
        return
    for cell_range in list(ws.merged_cells.ranges):
        if cell_range.min_row >= start_row:
            cell_range.shift(0, delta)


def _copy_xlsx_row_style(ws, source_row, target_row, max_column=9):
    from copy import copy
    if ws.row_dimensions[source_row].height is not None:
        ws.row_dimensions[target_row].height = ws.row_dimensions[source_row].height
    for column in range(1, max_column + 1):
        source = ws.cell(source_row, column)
        target = ws.cell(target_row, column)
        target._style = copy(source._style)
        target.protection = copy(source.protection)


def _ct_footer_rows(ws, start_row):
    wanted = {
        "subtotal": "tong truoc vat",
        "discount": "chiet khau dieu chinh",
        "vat": "tong tien vat",
        "total": "tong thanh toan",
    }
    found = {}
    for row in range(start_row, ws.max_row + 1):
        for column in range(1, 10):
            value = ws.cell(row, column).value
            if not isinstance(value, str):
                continue
            key = _normalise_docx_header(value)
            for name, label in wanted.items():
                if name not in found and key.startswith(label):
                    found[name] = row
    if set(found) != set(wanted):
        raise ValidationError("Template báo giá thiếu các dòng tổng bắt buộc.")
    return found


def _fill_ct_quote_xlsx(tpl_path, scalar_map, items, dynamic_contract=None):
    """Render the V3 one-sheet quotation contract with exactly N item rows."""
    import openpyxl
    from openpyxl.worksheet.properties import PageSetupProperties

    dynamic_contract = dynamic_contract or {}
    wb = openpyxl.load_workbook(tpl_path)
    if "BAO_GIA" not in wb.sheetnames:
        wb.close()
        raise ValidationError("Template báo giá thiếu sheet BAO_GIA.")
    sheet_name = dynamic_contract.get("sheet") or "BAO_GIA"
    if sheet_name not in wb.sheetnames:
        wb.close()
        raise ValidationError("Quotation template is missing its configured worksheet.")
    ws = wb[sheet_name]
    header_row = int(dynamic_contract.get("item_header_row") or 14)
    prototype_row = int(dynamic_contract.get("item_template_row") or 15)
    insert_at = prototype_row + 1
    rows = list(items or [])
    count = len(rows)
    if count > 1:
        delta = count - 1
        ws.insert_rows(insert_at, delta)
        _shift_merged_rows(ws, insert_at, delta)
    elif count == 0:
        ws.delete_rows(prototype_row, 1)
        _shift_merged_rows(ws, prototype_row, -1)

    last_item_row = prototype_row + count - 1
    for offset, item in enumerate(rows):
        row = prototype_row + offset
        _copy_xlsx_row_style(ws, prototype_row, row)
        values = _ct_quote_row(item, offset + 1)
        for column, value in enumerate(values, 1):
            ws.cell(row, column, value=value)
        ws.cell(row, 9, value='=IFERROR(F{0}*G{0}*(1+IF(H{0}="",0,H{0})),0)'.format(row))

    validation_last = last_item_row if count else prototype_row
    for validation in ws.data_validations.dataValidation:
        refs = str(validation.sqref)
        if "E15" in refs:
            validation.sqref = "E15:E%d" % validation_last
        elif "H15" in refs:
            validation.sqref = "H15:H%d" % validation_last

    footer = _ct_footer_rows(ws, prototype_row + count)
    if count:
        ws.cell(footer["subtotal"], 9,
                '=SUMPRODUCT(F15:F{0},G15:G{0})'.format(last_item_row))
        ws.cell(footer["vat"], 9,
                '=SUMPRODUCT(F15:F{0},G15:G{0},H15:H{0})'.format(last_item_row))
    else:
        ws.cell(footer["subtotal"], 9, "=0")
        ws.cell(footer["vat"], 9, "=0")
    ws.cell(footer["total"], 9, "=I%d-I%d+I%d" %
            (footer["subtotal"], footer["discount"], footer["vat"]))

    for row in ws.iter_rows():
        for cell in row:
            if isinstance(cell.value, str) and "{{" in cell.value:
                cell.value = _fill_text(cell.value, scalar_map, [], [0])

    ws.print_area = "A1:I%d" % ws.max_row
    ws.page_setup.orientation = "landscape"
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 0
    if ws.sheet_properties.pageSetUpPr is None:
        ws.sheet_properties.pageSetUpPr = PageSetupProperties(fitToPage=True)
    else:
        ws.sheet_properties.pageSetUpPr.fitToPage = True
    ws.print_title_rows = "%d:%d" % (header_row, header_row)
    if getattr(wb, "calculation", None) is not None:
        wb.calculation.calcMode = "auto"
        wb.calculation.fullCalcOnLoad = True
        wb.calculation.forceFullCalc = True
    buffer = io.BytesIO()
    wb.save(buffer)
    wb.close()
    return buffer.getvalue()


def _assert_ct_generation_supported(ma_mau, tpl_info):
    """Fail closed for V3.1 templates that lack structured rendering data.

    V3.1 has reusable prototype rows marked ``{{...}}``.  The legacy generic
    renderer erases unknown placeholders, therefore an unwired template must
    never be exported as a deceptively blank dossier.

    Allowed:
      - BG-* (XLSX dynamic quote rows from official quotation)
      - CT/HD with personnel / BOQ / journal / acceptance table wiring
      - CT-* and HD-* DOCX scalar fill from project + contract + quote leaf
        (static tables in the sample stay for manual completion)
    """
    if not str(tpl_info.get("registry_version") or "").startswith("3"):
        return
    if ma_mau.startswith("BG-"):
        return
    if (ma_mau in _CT_PERSONNEL_TABLES or ma_mau in _CT_BOQ_APPENDIX_TABLES
            or ma_mau in _CT_ACCEPTANCE_TABLES or ma_mau in _CT_JOURNAL_TABLES):
        return
    fmt = (tpl_info.get("format") or "").lower()
    if ma_mau.startswith(("CT-", "HD-")) and fmt in ("docx", "doc"):
        return
    raise ValidationError(
        "V3.1 template %s is registered but has no structured data wiring; "
        "export is blocked to prevent a blank document." % ma_mau)


def ct_auto_generation_status(ma_mau, tpl_info=None):
    """Expose a safe UI/API capability flag without attempting an export."""
    info = tpl_info or ct_templates().get(ma_mau)
    if not info:
        return False, "Template code is not registered."
    if ma_mau in _CT_ACCEPTANCE_TABLES:
        return False, "Mẫu đã wiring; sinh từ workspace Nghiệm thu để bắt buộc đủ gate và exact BOQ."
    if ma_mau in _CT_JOURNAL_TABLES:
        return False, "Sinh từ tab Nhật ký khi bản ghi đã duyệt (không xuất trống từ checklist)."
    try:
        _assert_ct_generation_supported(ma_mau, info)
    except ValidationError as exc:
        return False, str(exc)
    return True, ""


def _ct_versioned_filename(ma_mau, title, project_code, extension):
    safe_title = re.sub(r'[<>:"/\\|?*]', "", title or ma_mau)[:60].replace(" ", "")
    stamp = "%s_%s" % (datetime.now().strftime("%Y%m%d_%H%M%S_%f"), uuid.uuid4().hex[:8])
    return "%s_%s_%s_%s.%s" % (ma_mau, safe_title, project_code, stamp, extension)


def export_ct_doc(conn, sess, project_id, ma_mau, extra=None, defer_commit=False):
    """Sinh 1 DOCX/XLSX V2 cho project. Tra (fname, bytes, abs_path|None).
    Luu 1 ban vao folder khach ('Hồ sơ công trình') + index nhu moi chung tu khac.
    Filename luôn versioned để không ghi đè bản nháp/đã ký hiện hữu.
    KHONG dong cham trang_thai o day — api_write.ct_sinh_ho_so quan vong doi."""
    tpl_info = ct_templates().get(ma_mau)
    if not tpl_info:
        raise ValidationError("Mã mẫu không tồn tại: %s (xem TEMPLATE_MAPPING)" % ma_mau)
    tpl_path = tpl_info.get("abs_path") or _resolve_ct_template_path(tpl_info["path"])
    if not os.path.isfile(tpl_path):
        raise ValidationError("Thiếu file template: " + tpl_info["path"])
    project = conn.execute("SELECT * FROM project WHERE id=?", (project_id,)).fetchone()
    if not project:
        raise ValidationError("Công trình không tồn tại.")
    output_format = (tpl_info.get("format") or os.path.splitext(tpl_path)[1].lstrip(".")).lower()
    _assert_ct_generation_supported(ma_mau, tpl_info)
    if output_format == "xlsx":
        quotation, items = _ct_quotation_items(conn, project_id)
        if not quotation:
            raise ValidationError("Công trình chưa có báo giá để xuất template XLSX.")
        m = _ct_data_map(conn, project, sess, extra, quotation=quotation)
        data = _fill_ct_quote_xlsx(tpl_path, m, items, tpl_info.get("dynamic_rows"))
    elif output_format == "docx":
        personnel = _ct_project_personnel(conn, project_id) if ma_mau in _CT_PERSONNEL_TABLES else []
        dynamic_tables = _ct_personnel_dynamic_spec(ma_mau, personnel)
        profile_quotation = profile_contract = None
        journal_scalar = {}
        if ma_mau in _CT_JOURNAL_TABLES:
            _journal, journal_scalar, journal_specs = _ct_journal_doc_data(
                conn, project_id, (extra or {}).get("journal_id"))
            dynamic_tables += journal_specs
        if ma_mau in _CT_BOQ_APPENDIX_TABLES:
            active_profile, boq_rows = _ct_profile_boq_rows(conn, project_id)
            if not active_profile:
                raise ValidationError(
                    "Công trình chưa có BOQ chính thức active; không sinh phụ lục từ dự đoán cũ.")
            profile_quotation = conn.execute(
                "SELECT * FROM quotation WHERE id=? AND project_id=?",
                (_row_value(active_profile, "quotation_id"), project_id),
            ).fetchone()
            profile_contract = conn.execute(
                "SELECT * FROM hop_dong_ct WHERE id=? AND customer_id=?",
                (_row_value(active_profile, "contract_id"), project["customer_id"]),
            ).fetchone()
            if not profile_quotation or not profile_contract:
                raise ValidationError(
                    "Profile active thiếu liên kết báo giá/hợp đồng chính thức cho HD-07.")
            if int(_row_value(profile_contract, "quotation_id") or 0) != int(
                    _row_value(profile_quotation, "id") or 0):
                raise ValidationError(
                    "Liên kết hợp đồng và báo giá của profile active không đồng nhất.")
            dynamic_tables += _ct_boq_dynamic_spec(ma_mau, boq_rows)
        if ma_mau in _CT_ACCEPTANCE_TABLES:
            try:
                acceptance_id = int((extra or {}).get("acceptance_id"))
            except (TypeError, ValueError):
                raise ValidationError("Mẫu BBNT phải được sinh từ một đợt nghiệm thu hợp lệ.")
            _acceptance, acceptance_rows, acceptance_issues = _ct_acceptance_doc_rows(
                conn, project_id, acceptance_id)
            dynamic_tables += _ct_acceptance_dynamic_spec(
                ma_mau, acceptance_rows, acceptance_issues)
        m = _ct_data_map(conn, project, sess, extra,
                         quotation=profile_quotation, contract=profile_contract)
        m.update(journal_scalar)
        if ma_mau in _CT_BOQ_APPENDIX_TABLES:
            _validate_hd07_payment_schedule(m)
        data = _export_docx(tpl_path, m, [], dynamic_tables=dynamic_tables)
    else:
        raise ValidationError("Định dạng template hồ sơ chưa được hỗ trợ: " + output_format)
    fname = _ct_versioned_filename(ma_mau, tpl_info["title"],
                                    _row_value(project, "code"), output_format)
    abs_path = None
    try:
        is_personnel = ma_mau in _CT_PERSONNEL_TABLES
        generated_profile_role = (
            "personnel" if is_personnel else
            "contract" if ma_mau.startswith("HD-") else
            "official_quote" if ma_mau.startswith("BG-") else None)
        r = luu_file_vao_folder_khach(
            conn, project["customer_id"], "ho_so_cong_trinh_ct", fname, data,
            project_id=project_id, profile_role=generated_profile_role,
            commit=not defer_commit)
        if r.get("ok"):
            candidate_path = r["abs_path"]
            if is_personnel:
                tagged = conn.execute("""SELECT project_id,profile_role FROM source_document
                    WHERE abs_path=?""", (candidate_path,)).fetchone()
                if not tagged or tagged["project_id"] != project_id \
                        or tagged["profile_role"] != "personnel":
                    raise ValidationError(
                        "Không gắn được metadata công trình cho hồ sơ nhân sự.")
            abs_path = candidate_path
    except Exception:
        pass
    return fname, data, abs_path


def export_doc(conn, loai, doc_id, fmt="xlsx"):
    if loai not in DOC_TABLES:
        raise ValidationError("Loại chứng từ không hỗ trợ: " + loai)
    if fmt not in ("xlsx", "docx"):
        raise ValidationError("Định dạng phải là xlsx hoặc docx.")
    table, item_table, fk, title = DOC_TABLES[loai]
    doc = conn.execute("SELECT * FROM %s WHERE id=?" % table, (doc_id,)).fetchone()
    if not doc:
        raise ValidationError("Chứng từ không tồn tại.")
    items = []
    if item_table:
        items = [dict(r) for r in conn.execute(
            "SELECT * FROM %s WHERE %s=?" % (item_table, fk), (doc_id,)).fetchall()]
    # WO-29fix: 2 loai header-only van co bang dong trong TEMPLATE — cap dong cho chung:
    if not items and loai == "hop_dong" and ("quotation_id" in doc.keys()) and doc["quotation_id"]:
        # hop dong: keo hang muc tu bao gia goc ({{giai_doan_N}}/{{dvt_N}}/{{so_luong_N}}...)
        items = [dict(r) for r in conn.execute(
            "SELECT * FROM quotation_item WHERE quotation_id=? ORDER BY stt",
            (doc["quotation_id"],)).fetchall()]
    if not items and loai == "dccn":
        # dccn: 1 dong doi chieu tong hop tu chinh header ({{date_1}}/{{increase_1}}...)
        items = [{"date": fmt_d(date.today().isoformat()),
                  "description": "Phát sinh kỳ này (theo báo giá/quyết toán)",
                  "increase": fmt_vnd(doc["phat_sinh_tang"]), "payment": fmt_vnd(doc["da_thu"]),
                  "balance": fmt_vnd(doc["du_cuoi"]), "ghi_chu": ""}]
    cfg_row = conn.execute("SELECT * FROM cau_hinh WHERE id=1").fetchone()
    cfg = dict(cfg_row) if cfg_row else {}
    m = _data_map(conn, loai, doc, items, cfg)
    m["_title"] = title.upper()
    tpl = _find_tpl(loai, fmt)
    data = None
    if loai == "quotation" and fmt == "xlsx" and tpl:
        data = _fill_quote_template(tpl, m, items)  # khu VAT 10/8 + cong thuc tu tinh
    if data is None:
        data = _export_xlsx(None if (loai == "quotation" and fmt == "xlsx") else tpl,
                            m, items) if fmt == "xlsx" else _export_docx(tpl, m, items)
    kh_ten = m.get("customer_name", "").strip()[:40] or "KH"
    fname = "%s_%s_%s.%s" % (title.replace(" ", ""), re.sub(r'[<>:"/\\|?*]', "", kh_ten),
                             doc["code"], fmt)
    # Luu 1 ban vao dung thu muc con cua khach tren o D (best-effort, khong lam hong tai ve)
    try:
        cid = doc["customer_id"] if "customer_id" in doc.keys() else None
        if cid:
            luu_file_vao_folder_khach(conn, cid, loai, fname, data)
    except Exception:
        pass
    return fname, data
