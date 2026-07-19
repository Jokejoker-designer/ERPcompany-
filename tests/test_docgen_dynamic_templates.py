# -*- coding: utf-8 -*-
"""Isolated regression tests for dynamic CT personnel and quotation templates.

All records are synthetic.  Tests use in-memory SQLite and temporary Office files;
they never open the production DB and never write into a customer/project folder.
"""
import contextlib
import io
import os
import sqlite3
import tempfile
import unittest
from unittest import mock

import openpyxl
from docx import Document
from docx.shared import Pt
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.worksheet.datavalidation import DataValidation

import docgen as DG


APP_ROOT = os.path.dirname(os.path.abspath(__file__))
SCHEMA_PATH = os.path.join(APP_ROOT, "schema.sql")


DSNS_HEADERS = ("STT", "Họ và tên", "Chức vụ/Nhiệm vụ", "CCCD/CMND",
                "Năm sinh", "Số điện thoại", "Ghi chú")
PKBNV_HEADERS = ("STT", "Họ và tên", "Ngày/tháng/năm sinh", "CCCD",
                 "Số điện thoại", "Chức vụ", "Ghi chú")


def _make_conn(full_schema=False):
    conn = sqlite3.connect(":memory:")
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA foreign_keys=ON")
    if full_schema:
        with open(SCHEMA_PATH, encoding="utf-8") as handle:
            conn.executescript(handle.read())
    else:
        conn.executescript("""
            CREATE TABLE nhan_su(
                id INTEGER PRIMARY KEY, ho_ten TEXT, loai TEXT, sdt TEXT,
                cccd TEXT, ngay_sinh TEXT
            );
            CREATE TABLE project_personnel(
                id INTEGER PRIMARY KEY, project_id INTEGER, nhan_su_id INTEGER,
                source_row INTEGER, source_stt TEXT, project_role TEXT,
                site_role TEXT, source_note TEXT
            );
        """)
    return conn


def _build_personnel_template(path, headers):
    document = Document()
    document.add_paragraph("{{TEN_CONG_TRINH}} / {{CT}}")
    document.add_table(rows=1, cols=1)
    document.add_table(rows=1, cols=1)
    table = document.add_table(rows=4, cols=len(headers))
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    table.rows[1].height = Pt(24)
    for cell in table.rows[1].cells:
        run = cell.paragraphs[0].add_run("{{...}}")
        run.bold = True
    for row in table.rows[2:]:
        for cell in row.cells:
            cell.text = ""
    document.add_table(rows=1, cols=1)
    document.save(path)


def _build_quote_template(path):
    workbook = openpyxl.Workbook()
    ws = workbook.active
    ws.title = "BAO_GIA"
    ws["A1"] = "{{project_name}}"
    ws["A2"] = "{{customer_name}}"
    ws["A3"] = "{{quote_no}}"
    headers = ["STT", "MÃ / NHÓM", "TÊN HÀNG HÓA / HẠNG MỤC",
               "QUY CÁCH / HÃNG / MODEL", "ĐVT", "SỐ LƯỢNG",
               "ĐƠN GIÁ", "VAT %", "THÀNH TIỀN"]
    for column, value in enumerate(headers, 1):
        ws.cell(14, column, value)
    prototype = ["{{item_stt}}", "{{item_code_or_group}}", "{{item_name}}",
                 "{{item_specification}}", "{{item_uom}}", None, None, None,
                 '=IFERROR(F15*G15*(1+IF(H15="",0,H15)),0)']
    thin = Side(style="thin", color="000000")
    for column, value in enumerate(prototype, 1):
        cell = ws.cell(15, column, value)
        cell.fill = PatternFill("solid", fgColor="DDEBF7")
        cell.font = Font(name="Arial", bold=True)
        cell.alignment = Alignment(vertical="center", wrap_text=True)
        cell.border = Border(left=thin, right=thin, top=thin, bottom=thin)
    ws["F15"].number_format = '#,##0.###'
    ws["G15"].number_format = '#,##0 "đ"'
    ws["H15"].number_format = "0%"
    ws["I15"].number_format = '#,##0 "đ"'
    ws.row_dimensions[15].height = 24
    footer = [
        (17, "TỔNG TRƯỚC VAT", '=SUMPRODUCT(F15:F15,G15:G15)'),
        (18, "CHIẾT KHẤU / ĐIỀU CHỈNH", 0),
        (19, "TỔNG TIỀN VAT", '=SUMPRODUCT(F15:F15,G15:G15,H15:H15)'),
        (20, "TỔNG THANH TOÁN", "=I17-I18+I19"),
    ]
    for row, label, formula in footer:
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=5)
        ws.merge_cells(start_row=row, start_column=6, end_row=row, end_column=8)
        ws.cell(row, 6, label)
        ws.cell(row, 9, formula)
    ws.merge_cells("A21:I21")
    ws["A21"] = "Bằng chữ: {{amount_in_words}}"
    ws.merge_cells("A31:I31")
    ws["A31"] = "END"
    uom = DataValidation(type="list", formula1='"bộ,cái,m"')
    vat = DataValidation(type="list", formula1='"0%,8%,10%"')
    ws.add_data_validation(uom)
    ws.add_data_validation(vat)
    uom.add(ws["E15"])
    vat.add(ws["H15"])
    workbook.save(path)
    workbook.close()


def _footer_rows(ws):
    out = {}
    labels = {
        "TỔNG TRƯỚC VAT": "subtotal",
        "CHIẾT KHẤU / ĐIỀU CHỈNH": "discount",
        "TỔNG TIỀN VAT": "vat",
        "TỔNG THANH TOÁN": "total",
    }
    for row in range(1, ws.max_row + 1):
        for cell in ws[row]:
            if cell.value in labels:
                out[labels[cell.value]] = row
    return out


class PersonnelDynamicTableTest(unittest.TestCase):
    def setUp(self):
        self.conn = _make_conn()
        self.tmp = tempfile.TemporaryDirectory()
        self.conn.executemany(
            "INSERT INTO nhan_su(id,ho_ten,loai,sdt,cccd,ngay_sinh) VALUES(?,?,?,?,?,?)",
            [(1, "Nhan Su A", "KTV", "0900000001", "000000000001", "1990-02-03"),
             (2, "Nhan Su B", "Tho", "0900000002", "000000000002", "1988"),
             (3, "Unassigned", "KTV", "0900000003", "000000000003", "1995")])
        self.conn.executemany("""INSERT INTO project_personnel
            (id,project_id,nhan_su_id,source_row,source_stt,project_role,site_role,source_note)
            VALUES(?,?,?,?,?,?,?,?)""",
            [(1, 7, 1, 3, "01", "Chỉ huy", "KTV", ""),
             (2, 7, 2, 4, "02", "", "Công nhân", "")])

    def tearDown(self):
        self.conn.close()
        self.tmp.cleanup()

    def test_assignments_only_and_both_personnel_templates_preserve_style(self):
        stdout, stderr = io.StringIO(), io.StringIO()
        with contextlib.redirect_stdout(stdout), contextlib.redirect_stderr(stderr):
            rows = DG._ct_project_personnel(self.conn, 7)
        self.assertEqual(stdout.getvalue(), "")
        self.assertEqual(stderr.getvalue(), "")
        self.assertEqual(len(rows), 2)
        self.assertEqual(rows[0]["nam_sinh"], "1990")
        self.assertEqual(rows[0]["ngay_sinh"], "03/02/1990")

        for code, headers in (("CT-01-DSNS", DSNS_HEADERS),
                              ("CT-01-PKBNV", PKBNV_HEADERS)):
            with self.subTest(code=code):
                path = os.path.join(self.tmp.name, code + ".docx")
                _build_personnel_template(path, headers)
                data = DG._export_docx(
                    path, {"TEN_CONG_TRINH": "Project Test", "CT": "Project Test"}, [],
                    dynamic_tables=DG._ct_personnel_dynamic_spec(code, rows))
                rendered = Document(io.BytesIO(data))
                table = rendered.tables[2]
                self.assertEqual(len(table.rows), 3)  # header + exactly two assignments
                self.assertEqual(table.rows[1].cells[1].text, "Nhan Su A")
                self.assertEqual(table.rows[2].cells[1].text, "Nhan Su B")
                self.assertTrue(table.rows[1].cells[0].paragraphs[0].runs[0].bold)
                self.assertNotIn("{{", "\n".join(
                    cell.text for tab in rendered.tables for row in tab.rows for cell in row.cells))


class DynamicQuotationWorkbookTest(unittest.TestCase):
    def setUp(self):
        self.tmp = tempfile.TemporaryDirectory()
        self.template = os.path.join(self.tmp.name, "quote.xlsx")
        _build_quote_template(self.template)
        self.scalars = {
            "project_name": "Project Test", "customer_name": "Customer Test",
            "quote_no": "BG-TEST", "amount_in_words": "Một trăm đồng",
        }

    def tearDown(self):
        self.tmp.cleanup()

    def _items(self, count):
        return [{
            "source_stt": str(index), "item_code_or_group": "G-%d" % index,
            "hang_muc": "Hang muc %d" % index, "technical_requirement": "Spec",
            "dvt": "bộ", "so_luong": index + 0.5, "don_gia": index * 1000,
            "thue_suat": 8 if index % 2 else "10%",
        } for index in range(1, count + 1)]

    def test_exact_item_counts_formulas_styles_validations_and_print_setup(self):
        source_row = DG._ct_quote_row({
            "source_stt": "A.1", "hang_muc": "Source item", "contract_qty_raw": "2,5 bộ",
            "so_luong": 0, "unit_price_raw": "1.234.000", "don_gia": 0,
            "vat_rate_raw": "8%", "thue_suat": 0,
        }, 1)
        self.assertEqual(source_row[5], 2.5)
        self.assertEqual(source_row[6], 1234000.0)
        self.assertEqual(source_row[7], 0.08)
        for count in (1, 3, 10, 31, 75, 150):
            with self.subTest(count=count):
                data = DG._fill_ct_quote_xlsx(self.template, self.scalars, self._items(count))
                workbook = openpyxl.load_workbook(io.BytesIO(data), data_only=False)
                ws = workbook["BAO_GIA"]
                last_item = 14 + count
                self.assertEqual([ws.cell(row, 3).value for row in range(15, last_item + 1)],
                                 ["Hang muc %d" % i for i in range(1, count + 1)])
                self.assertEqual(ws["F15"].value, 1.5)
                self.assertEqual(ws["G15"].value, 1000.0)
                self.assertEqual(ws["H15"].value, 0.08)
                self.assertEqual(ws["I15"].value,
                                 '=IFERROR(F15*G15*(1+IF(H15="",0,H15)),0)')
                self.assertEqual(ws.cell(last_item, 1)._style, ws["A15"]._style)
                footer = _footer_rows(ws)
                self.assertEqual(ws.cell(footer["subtotal"], 9).value,
                                 '=SUMPRODUCT(F15:F{0},G15:G{0})'.format(last_item))
                self.assertEqual(ws.cell(footer["vat"], 9).value,
                                 '=SUMPRODUCT(F15:F{0},G15:G{0},H15:H{0})'.format(last_item))
                self.assertEqual(ws.cell(footer["total"], 9).value,
                                 "=I%d-I%d+I%d" % (footer["subtotal"], footer["discount"], footer["vat"]))
                refs = {str(validation.sqref) for validation in ws.data_validations.dataValidation}
                expected_e = "E15" if count == 1 else "E15:E%d" % last_item
                expected_h = "H15" if count == 1 else "H15:H%d" % last_item
                self.assertIn(expected_e, refs)
                self.assertIn(expected_h, refs)
                self.assertEqual(str(ws.print_area), "'BAO_GIA'!$A$1:$I$%d" % ws.max_row)
                self.assertEqual(ws.print_title_rows, "$14:$14")
                self.assertEqual(ws.page_setup.orientation, "landscape")
                self.assertEqual(ws.page_setup.fitToHeight, 0)
                self.assertNotIn("{{", "\n".join(
                    str(cell.value) for row in ws.iter_rows() for cell in row if cell.value is not None))
                workbook.close()


class ExportCtDocIntegrationTest(unittest.TestCase):
    def setUp(self):
        self.conn = _make_conn(full_schema=True)
        self.tmp = tempfile.TemporaryDirectory()
        self.personnel_template = os.path.join(self.tmp.name, "personnel.docx")
        self.quote_template = os.path.join(self.tmp.name, "quote.xlsx")
        _build_personnel_template(self.personnel_template, DSNS_HEADERS)
        _build_quote_template(self.quote_template)
        self.conn.execute("""INSERT INTO customer
            (id,code,customer_name,phan_loai,tax_id,dia_chi,nguoi_lien_he)
            VALUES(1,'KH-TEST','Customer Test','Cong ty','TAX','Address','Contact')""")
        self.conn.execute("""INSERT INTO project
            (id,code,project_name,customer_id,status,dia_diem)
            VALUES(7,'CT-TEST','Project Test',1,'Working','Location')""")
        self.conn.execute("""INSERT INTO cau_hinh(id,ten_cong_ty,ma_so_thue,dia_chi)
                           VALUES(1,'Contractor Test','TAX-B','Address B')""")
        self.conn.execute("""INSERT INTO quotation
            (id,code,customer_id,project_id,grand_total,status,is_official,ngay_lap)
            VALUES(9,'BG-TEST',1,7,2200,'Da duyet',1,'2026-07-13')""")
        self.conn.executemany("""INSERT INTO quotation_item
            (id,quotation_id,stt,hang_muc,khoi_luong,don_gia,thanh_tien,
             technical_requirement,source_stt,source_item_raw,source_row)
            VALUES(?,?,?,?,?,?,?,?,?,?,?)""",
            [(1, 9, 1, "Item A", "2 bộ", 1000, 2000, "Spec A", "1", "Item A", 20),
             (2, 9, 2, "Item B", "1 cái", 200, 200, "Spec B", "2", "Item B", 21)])
        self.conn.execute("""INSERT INTO nhan_su(id,ho_ten,loai,sdt,cccd,ngay_sinh)
                           VALUES(1,'Nhan Su A','KTV','0900000001','000000000001','1990-02-03')""")
        self.conn.execute("""INSERT INTO project_personnel
            (id,project_id,nhan_su_id,source_row,source_stt,project_role)
            VALUES(1,7,1,2,'1','Chỉ huy')""")
        self.conn.commit()

    def tearDown(self):
        self.conn.close()
        self.tmp.cleanup()

    def test_export_routes_docx_and_xlsx_without_overwrite_or_pii_logging(self):
        registry = {
            "CT-01-DSNS": {"title": "Danh sach nhan su", "path": self.personnel_template,
                            "abs_path": self.personnel_template, "format": "docx"},
            "BG-01_BAO_GIA_DA_NANG": {"title": "Bao gia", "path": self.quote_template,
                                       "abs_path": self.quote_template, "format": "xlsx"},
        }
        stdout, stderr = io.StringIO(), io.StringIO()
        with mock.patch.object(DG, "ct_templates", return_value=registry), \
                mock.patch.object(DG, "luu_file_vao_folder_khach", return_value={"ok": False}), \
                contextlib.redirect_stdout(stdout), contextlib.redirect_stderr(stderr):
            doc_name, doc_data, doc_path = DG.export_ct_doc(
                self.conn, {"full_name": "Author Test"}, 7, "CT-01-DSNS")
            xlsx_name, xlsx_data, xlsx_path = DG.export_ct_doc(
                self.conn, {"full_name": "Author Test"}, 7, "BG-01_BAO_GIA_DA_NANG")
        self.assertEqual(stdout.getvalue(), "")
        self.assertEqual(stderr.getvalue(), "")
        self.assertIsNone(doc_path)
        self.assertIsNone(xlsx_path)
        project = self.conn.execute("SELECT * FROM project WHERE id=7").fetchone()
        aliases = DG._ct_data_map(self.conn, project, {"full_name": "Author Test"})
        self.assertEqual(aliases["CT"], "Project Test")
        self.assertEqual(aliases["DD"], "Location")
        self.assertEqual(aliases["CDT"], "Customer Test")
        self.assertEqual(aliases["quote_no"], "BG-TEST")
        self.assertEqual(aliases["customer_tax_id"], "TAX")
        self.assertRegex(doc_name, r"_\d{8}_\d{6}_\d{6}_[0-9a-f]{8}\.docx$")
        self.assertRegex(xlsx_name, r"_\d{8}_\d{6}_\d{6}_[0-9a-f]{8}\.xlsx$")
        self.assertNotEqual(
            DG._ct_versioned_filename("CT-01-DSNS", "Danh sach", "CT-TEST", "docx"),
            DG._ct_versioned_filename("CT-01-DSNS", "Danh sach", "CT-TEST", "docx"),
        )
        rendered_doc = Document(io.BytesIO(doc_data))
        self.assertEqual(rendered_doc.tables[2].rows[1].cells[1].text, "Nhan Su A")
        rendered_xlsx = openpyxl.load_workbook(io.BytesIO(xlsx_data), data_only=False)
        ws = rendered_xlsx["BAO_GIA"]
        self.assertEqual(ws["A1"].value, "Project Test")
        self.assertEqual(ws["A2"].value, "Customer Test")
        self.assertEqual(ws["F15"].value, 2.0)
        self.assertEqual(ws["G15"].value, 1000.0)
        rendered_xlsx.close()

    def test_personnel_export_tags_source_document_with_project_scope(self):
        registry = {
            "CT-01-DSNS": {"title": "Danh sach nhan su", "path": self.personnel_template,
                            "abs_path": self.personnel_template, "format": "docx"},
        }

        def save_and_index(conn, customer_id, _doc_type, filename, _data,
                           project_id=None, profile_role=None, commit=True):
            path = os.path.join(self.tmp.name, filename)
            conn.execute("""INSERT INTO source_document(customer_id,project_id,profile_role,
                            file_name,rel_path,abs_path,ext) VALUES(?,?,?,?,?,? ,'.docx')""",
                         (customer_id, project_id, profile_role, filename,
                          os.path.join("test", filename), path))
            if commit:
                conn.commit()  # Match luu_file_vao_folder_khach's existing contract.
            return {"ok": True, "abs_path": path}

        with mock.patch.object(DG, "ct_templates", return_value=registry), \
                mock.patch.object(DG, "luu_file_vao_folder_khach", side_effect=save_and_index):
            _name, _data, path = DG.export_ct_doc(
                self.conn, {"full_name": "Author Test"}, 7, "CT-01-DSNS")
        self.assertIsNotNone(path)
        self.conn.commit()  # Export callers commit metadata together with dossier status.
        indexed = self.conn.execute(
            "SELECT project_id,profile_role FROM source_document WHERE abs_path=?", (path,)
        ).fetchone()
        self.assertEqual((indexed["project_id"], indexed["profile_role"]), (7, "personnel"))

    def test_personnel_export_returns_no_path_when_index_metadata_cannot_be_tagged(self):
        registry = {
            "CT-01-DSNS": {"title": "Danh sach nhan su", "path": self.personnel_template,
                            "abs_path": self.personnel_template, "format": "docx"},
        }
        missing_index_path = os.path.join(self.tmp.name, "not-indexed.docx")
        with mock.patch.object(DG, "ct_templates", return_value=registry), \
                mock.patch.object(DG, "luu_file_vao_folder_khach",
                                  return_value={"ok": True, "abs_path": missing_index_path}):
            _name, _data, path = DG.export_ct_doc(
                self.conn, {"full_name": "Author Test"}, 7, "CT-01-DSNS")
        self.assertIsNone(path)


class V31RegistrySafetyTest(unittest.TestCase):
    def test_v31_registry_and_profile_rules_are_complete(self):
        previous_registry = DG._CT_TPL_CACHE
        previous_rules = DG._CT_REQUIRED_CACHE
        try:
            DG._CT_TPL_CACHE = None
            DG._CT_REQUIRED_CACHE = None
            registry = DG.ct_templates()
            rules = DG.ct_document_requirements("INSTALLATION_STANDARD")
            self.assertEqual(len(registry), 84)
            self.assertEqual(sum(row["format"] == "docx" for row in registry.values()), 69)
            self.assertEqual(sum(row["format"] == "xlsx" for row in registry.values()), 15)
            self.assertFalse(any(not row["exists"] for row in registry.values()))
            self.assertEqual(len(rules["required"]), 21)
            self.assertEqual(len(rules["conditional"]), 20)
            self.assertEqual(registry["HD-07"]["legacy_code"],
                             "HD-07_PHU_LUC_BANG_KHOI_LUO")
            self.assertEqual(registry["BG-03"]["dynamic_rows"]["item_template_row"], 15)
        finally:
            DG._CT_TPL_CACHE = previous_registry
            DG._CT_REQUIRED_CACHE = previous_rules

    def test_unwired_v31_template_is_blocked_but_wired_ones_are_allowed(self):
        registered = {"registry_version": "3.1"}
        with self.assertRaises(DG.ValidationError):
            DG._assert_ct_generation_supported("CT-05-BCTD", registered)
        for code in ("BG-03", "CT-01-DSNS", "CT-01-PKBNV", "HD-07"):
            with self.subTest(code=code):
                DG._assert_ct_generation_supported(code, registered)


if __name__ == "__main__":
    unittest.main()
